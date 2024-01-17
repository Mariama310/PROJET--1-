#imports
import pandas as pd
from PIL import Image, ImageTk
import tkinter as tk
from tkinter import ttk, simpledialog, messagebox
import re
import customtkinter as ctk
from datetime import date
import csv

#   pip install python-docx (for editing words)
from docx import Document
from docx.shared import Pt, RGBColor #text size, color
from docx.enum.text import WD_ALIGN_PARAGRAPH #paragraph alignment

from tkcalendar import Calendar #pip install tkcalendar


ctk.set_appearance_mode("Dark")  # Modes: "System" (standard), "Dark", "Light"
ctk.set_default_color_theme("blue")

class Personne:
    def __init__(self, personne_id, name=None, address=None, phone_number=None):
        self._personne_id = personne_id
        self._name = name
        self._address = address
        self._phone_number = phone_number
    
    def get_id(self):
        return self._personne_id
    
    def set_id(self, personne_id):
        self._personne_id = personne_id
    id=property(get_id,set_id)
    
    def get_name(self):
        return self._name
    def set_name(self, name):
        self._name = name
    name=property(get_name,set_name)
    
    def get_address(self):
        return self._address
    
    def set_address(self, address):
        self._address = address
    address=property(get_address,set_address)
    
    def get_phone_number(self):
        return self._phone_number
    
    def set_phone_number(self, phone_number):
        self._phone_number = phone_number
    phone=property(get_phone_number,set_phone_number)

class Supplier(Personne):
    
    def __init__(self, supplier_id, name=None, address=None, phone_number=None):
        super().__init__(supplier_id, name, address, phone_number)
        self._supplier_id = supplier_id
    
    def get_supplier_id(self):
        return self._supplier_id
    
    def set_supplier_id(self, supplier_id):
        self._supplier_id = supplier_id
    id=property(get_supplier_id,set_supplier_id)

class Employee(Personne):
    def __init__(self, employee_id, name=None, address=None, phone_number=None, position=None, salary=None, authority=None):
        super().__init__(employee_id, name, address, phone_number)
        self._employee_id = employee_id
        self._position = position
        self._salary = salary
        self.phone_number = phone_number
        self.name = name
        self._authority = authority

    def get_employee_id(self):
        return self._employee_id

    def set_employee_id(self, employee_id):
        self._employee_id = employee_id
    emp_id = property(get_employee_id, set_employee_id)

    def get_position(self):
        return self._position

    def set_position(self, position):
        self._position = position
    position = property(get_position, set_position)

    def get_salary(self):
        return self._salary

    def set_salary(self, salary):
        self._salary = salary
    salary = property(get_salary, set_salary)

    def get_phone_number(self):
        return self.phone_number

    def set_phone_number(self, phone):
        self.phone_number = phone
    phone = property(get_phone_number, set_phone_number)

    def get_authority(self):
        return self._authority

    def set_authority(self, authority):
        self._authority = authority
    authority = property(get_authority, set_authority)

class Client(Personne):
    def __init__(self, client_id, name=None, address=None, phone_number=None):
        super().__init__(client_id, name, address, phone_number)
        self._client_id = client_id
    def get_client_id(self):
        return self._client_id
    
    def set_client_id(self, client_id):
        self._client_id = client_id
    clt_id=property(get_client_id,set_client_id)

class Database:
    def __init__(self):
        self.employees = []
        self.products = []
        self.orders = []
        self.clients = []
        self.warehouse = Warehouse(capacity=1000) #on peut changer la capacité si nécessaire 
        self.sales = []
        self.suppliers = []
        self.invoices = []
        self.reports = []

    def add_employee(self, employee):
        self.employees.append(employee)

    def add_product(self, product):
        self.products.append(product)

    def add_order(self, order):
        self.orders.append(order)

    def add_client(self, client):
        self.clients.append(client)

    def add_sale(self, sale):
        self.sales.append(sale)

    def add_supplier(self, supplier):
        self.suppliers.append(supplier)

    def add_invoice(self, invoice):
        self.invoices.append(invoice)

    def add_report(self, report):
        self.reports.append(report)

    def show_database(self):
        print("Employees:")
        for employee in self.employees:
            print(employee.get_employee_id(), employee.get_name(), employee.get_position())
        
        print("\nProducts:")
        for product in self.products:
            print(product.get_product_id(), product.get_name(), product.get_price())
        
        print("\nOrders:")
        for order in self.orders:
            print(order.get_order_id(), order.get_client_id(), order.get_products(), order.get_order_date())
        
        print("\nClients:")
        for client in self.clients:
            print(client.get_client_id(), client.get_name(), client.get_address())
        
        print("\nSales:")
        for sale in self.sales:
            print(sale.get_sale_id(), sale.get_product_id(), sale.get_quantity(), sale.get_sale_date())
        
        print("\nSuppliers:")
        for supplier in self.suppliers:
            print(supplier.get_supplier_id(), supplier.get_name(), supplier.get_address())
        
        print("\nInvoices:")
        for invoice in self.invoices:
            print(invoice.get_invoice_id(), invoice.get_client_id(), invoice.get_invoice_date(), invoice.get_total_amount())
        
        print("\nReports:")
        for report in self.reports:
            print(report.get_report_id(), report.get_title(), report.get_report_date(), report.get_data())

class Product:
    def __init__(self, product_id, description=None, price=None, quantity_in_stock=None, historique = None):
        self._product_id = product_id
        self._description = description
        self._price = price
        self._quantity_in_stock = quantity_in_stock
        self.historique = historique
    
    def get_product_id(self):
        return self._product_id
    
    def set_product_id(self, product_id):
        self._product_id = product_id
    id=property(get_product_id,set_product_id)
    
    def get_description(self):
        return self._description
    
    def set_description(self, description):
        self._description = description
    description=property(get_description,set_description)

    def get_price(self):
        return self._price
    
    def set_price(self, price):
        self._price = price
    prix=property(get_price,set_price)
    
    def get_quantity_in_stock(self):
        return self._quantity_in_stock
    
    def set_quantity_in_stock(self, quantity_in_stock):
        self._quantity_in_stock = quantity_in_stock

    stock=property(get_quantity_in_stock,set_quantity_in_stock)
    
class Order:
    def __init__(self, order_id, order_date, client, products, payment_type="", price=0, price_paid = 0, pompe=False):
        self._order_id = order_id
        self._order_date = order_date
        self._client = client
        self._products = products
        self._price = price
        self._price_paid = price_paid
        self._payment_type=payment_type.capitalize()
        self._pompe=pompe
    
    def get_order_id(self):
        return self._order_id
    def set_order_id(self, order_id):
        self._order_id = order_id
    order_id=property(get_order_id,set_order_id)
    
    def get_order_date(self):
        return self._order_date
    def set_order_date(self, order_date):
        self._order_date = order_date
    order_date=property(get_order_date,set_order_date)
    
    def get_client(self):
        return self._client
    def set_client(self, client):
        self._client = client
    client=property(get_client,set_client)
    
    def get_products(self):
        return self._products
    def set_products(self, products):
        self._products = products
    products=property(get_products,set_products)
    
    def get_price(self):
        return self._price
    def set_price(self, price):
        self._price=price
    price=property(get_price, set_price)
    
    def get_price_paid(self):
        return self._price_paid
    def set_price_paid(self, price_paid):
        self._price_paid=price_paid
    price_paid=property(get_price_paid, set_price_paid)
    
    def get_payment_type(self):
        return self._payment_type
    def set_payment_type(self, payment_type):
        self._payment_type=payment_type.capitalize()
    payment_type=property(get_payment_type,set_payment_type)
    
    def get_pompe(self):
        return self._pompe
    def set_pompe(self, pompe):
        self._pompe=pompe
    pompe=property(get_pompe, set_pompe)
    
    def CreateBDC(self):
        titleSize = Pt(13)
        valueSize = Pt(12)
        tickSize = Pt(14)
        tickChar = '\u2714'
        tickColor = '2ec92e' #code hex sans le #80ff80
            
        def AddTitleValue(paragraph, title, value="", length=0, alignment='left'):
            if alignment=='right':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment=='center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = paragraph.add_run(title)
            run.bold=True
            run.font.size=titleSize
            paragraph.add_run(str(value)).font.size = valueSize
            if len(str(value))<length and len(str(value))!=0:
                paragraph.add_run(" "*(length-len(str(value)))).font.size = valueSize
        
        document = Document('BON DE COMMANDE - blank.docx')
        
        #date
        AddTitleValue(document.paragraphs[-1], 'CM DU : ', date.today().strftime("%d/%m/%Y"), alignment='right')
        
        #id commande
        AddTitleValue(document.add_paragraph(), 'BON DE COMMANDE N° : ', self.order_id, 28)
        
        #nom client
        p = document.add_paragraph()
        AddTitleValue(p, 'CLIENT : ', self.client.name, 27)
        
        #tel client
        AddTitleValue(p, '\t\t\t\t\t\tTEL : ', self.client.phone, 24)
        
        #products
        AddTitleValue(document.add_paragraph(), "Produit :\t\t\tQuantité :")
        for product in self.products:
            p = document.add_paragraph()
            AddTitleValue(p, "", product[0].description, 60)
            AddTitleValue(p, "\t\t\t", product[1])
        
        # dosage tick option
        p = document.add_paragraph()
        AddTitleValue(p, 'Dosage :\t')
        AddTitleValue(p, 'Pompé  ')
        tick1, tick2 = (tickChar, ' ') if self.pompe else (' ', tickChar)
        run = p.add_run(f"{tick1}\t")
        run.font.size = tickSize
        run.font.color.rgb = RGBColor.from_string(tickColor)
        AddTitleValue(p, 'Non Pompé  ')
        run = p.add_run(f"{tick2}")
        run.font.size = tickSize
        run.font.color.rgb = RGBColor.from_string(tickColor)
        
        #payment method
        if self.payment_type=='Especes':
            AddTitleValue(p, '\t\t\tRèglement : ','Espèces')
        elif self.payment_type=='Cheque':
            AddTitleValue(p, '\t\t\tRèglement : ','Chèque')
        else:
            AddTitleValue(p, '\t\t\tRèglement : ','Virement')

        #date commande
        p = document.add_paragraph()
        AddTitleValue(p, 'Date de livraison : ', self.order_date, 24)

        #adresse client
        AddTitleValue(p, '\t\t\t\tDestination : ', self.client.address, 28)
        
        #montant
        p = document.add_paragraph()
        AddTitleValue(p, 'Montant : ', self.price, 19)
        
        #reste
        AddTitleValue(p, '\t\t\t\t\t\tReste à payer : ', self.price-self.price_paid, 19)
        
        #signatures
        AddTitleValue(document.add_paragraph(), 'Signature du client :\t\t\t\t\t\t\t\t\tSignature du Chef de la Centrale : ', "")
        
        document.save(f'Bon de Commande - {self.order_id} .docx')

class Warehouse:
    def __init__(self, capacity=None,quantity_in_stock=None,products_in_stock=None):
        self._capacity = capacity
        self._products_in_stock = products_in_stock
        self._quantity_in_stock=quantity_in_stock

    def get_capacity(self):
        return self._capacity
    
    def set_capacity(self, capacity):
        self._capacity = capacity
    capacity=property(get_capacity,set_capacity)
    
    def get_products_in_stock(self):
        return self._products_in_stock
    
    def add_product(self, product):
        self._products_in_stock.append(product)
    
    def remove_product(self, product):
        self._products_in_stock.remove(product)
    products=property(get_products_in_stock,add_product,remove_product)

    def get_quantity_in_stock(self):
        return self._quantity_in_stock
    def set_quantity_in_stock(self,q):
        self._quantity_in_stock=q
    quantity=property(get_quantity_in_stock,set_quantity_in_stock)

class Warehouse:
    def __init__(self, capacity=None,quantity_in_stock=None,products_in_stock=None):
        self._capacity = capacity
        self._products_in_stock = products_in_stock
        self._quantity_in_stock=quantity_in_stock

    def get_capacity(self):
        return self._capacity
    
    def set_capacity(self, capacity):
        self._capacity = capacity
    capacity=property(get_capacity,set_capacity)
    
    def get_products_in_stock(self):
        return self._products_in_stock
    
    def add_product(self, product):
        self._products_in_stock.append(product)
    
    def remove_product(self, product):
        self._products_in_stock.remove(product)
    products=property(get_products_in_stock,add_product,remove_product)

    def get_quantity_in_stock(self):
        return self._quantity_in_stock
    def set_quantity_in_stock(self,q):
        self._quantity_in_stock=q
    quantity=property(get_quantity_in_stock,set_quantity_in_stock)

class Sale:
    def __init__(self,client,sale_id=None, sale_date=None, product=None, quantity_sold=None, sale_price=None):
        self._sale_date = sale_date
        self._product = product
        self._quantity_sold = quantity_sold
        self._sale_price = sale_price
        self._client = client

    def get_sale_id(self):
        return self._sale_id
    
    def set_sale_id(self, sale_id):
        self._sale_id = sale_id
    sale_id=property(get_sale_id,set_sale_id)
    
    def get_sale_date(self):
        return self._sale_date
    
    def set_sale_date(self, sale_date):
        self._sale_date = sale_date
    sale_date=property(get_sale_date,set_sale_date)
    
    def get_product(self):
        return self._product
    
    def set_product(self, product):
        self._product = product
    prod=property(get_product,set_product)
    
    def get_quantity_sold(self):
        return self._quantity_sold
    
    def set_quantity_sold(self, quantity_sold):
        self._quantity_sold = quantity_sold
    sold=property(get_quantity_sold,set_quantity_sold)
    
    def get_sale_price(self):
        return self._sale_price
    
    def set_sale_price(self, sale_price):
        self._sale_price = sale_price
    price=property(get_sale_price,set_sale_price)
    
    def get_client(self):
        return self._client
    
    def set_client(self, client):
        self._client = client
    client=property(get_client,set_client)

    def report(self, products):
        sales_per_client = {}
        for sale in self.sales:
            client = sale.get_client().get_client_id()
            quantity_sold = sale.get_quantity_sold()
            if client in sales_per_client:
                sales_per_client[client] += quantity_sold
            else:
                sales_per_client[client] = quantity_sold

        average_sales_per_client = {client: total_sales / len(self.sales) for client, total_sales in sales_per_client.items()}

        sales_per_product = {}
        for sale in self.sales:
            product = sale.get_product().get_product_id()
            quantity_sold = sale.get_quantity_sold()
            if product in sales_per_product:
                sales_per_product[product] += quantity_sold
            else:
                sales_per_product[product] = quantity_sold

        total_revenue = sum(sale.get_sale_price() * sale.get_quantity_sold() for sale in self.sales)

        delivery_times = [(sale.get_sale_date() - sale.get_client().get_registration_date()).days for sale in self.sales]
        average_delivery_time = sum(delivery_times) / len(delivery_times)

        most_demanded_product = max(sales_per_product, key=sales_per_product.get)

        return {
            "average_sales_per_client": average_sales_per_client,
            "total_sales_per_product": sales_per_product,
            "total_revenue": total_revenue,
            "average_delivery_time": average_delivery_time,
            "most_demanded_product": most_demanded_product }

  
    def generate_report(self):
        revenue = self._quantity_sold * self._sale_price
        report_date = self._sale_date 
        report = report(report_date, revenue)
        return report
    
class Invoice:
    def __init__(self, invoice_id, invoice_date=None, total_amount=None, products=None):
        self._invoice_id = invoice_id
        self._invoice_date = invoice_date
        self._total_amount = total_amount
        self._products = products
    
    def get_invoice_id(self):
        return self._invoice_id
    
    def set_invoice_id(self, invoice_id):
        self._invoice_id = invoice_id
    id=property(get_invoice_id,set_invoice_id)
    
    def get_invoice_date(self):
        return self._invoice_date
    
    def set_invoice_date(self, invoice_date):
        self._invoice_date = invoice_date
    date=property(get_invoice_date,set_invoice_date)
    
    def get_total_amount(self):
        return self._total_amount
    
    def set_total_amount(self, total_amount):
        self._total_amount = total_amount
    total_amount=property(get_total_amount,set_total_amount)
    
    def get_products(self):
        return self._products
    
    def set_products(self, products):
        self._products = products
    products=property(get_products,set_products)

class Livraison:
    def __init__(self, shipping_nbr, client_name, cin, shipping_address, phone_nbr, order_nbr, order_date, 
                 central_dep_hr, central_arr_hr, worksite_arr_hr, worksite_dep_hr, product_id, quantity, 
                 product_type, vehicle, total_ttc):
        self._shipping_nbr = shipping_nbr
        self._client_name = client_name
        self._cin = cin
        self._shipping_address = shipping_address
        self._phone_nbr = phone_nbr
        self._order_nbr = order_nbr
        self._order_date = order_date
        self._central_dep_hr = central_dep_hr
        self._central_arr_hr = central_arr_hr
        self._worksite_arr_hr = worksite_arr_hr # retirer les heures d'arriver et de départ du lieu de livraison car il faut simplement le temps total qu'il a fallu mobiliser un vehicule
        self._worksite_dep_hr = worksite_dep_hr #
        self._product_id = product_id
        self._quantity = quantity
        self._product_type = product_type # à retirer car accessible depuis product_id
        self._vehicle = vehicle
        self._total_ttc = total_ttc # à retirer plus facil à calculer que de stocker
        
    def __init__(self, shipping_nbr, client_name, shipping_address, phone_nbr, order_nbr, order_date, 
                 central_dep_hr, central_arr_hr, worksite_arr_hr, worksite_dep_hr, product_id, quantity, 
                 product_type, vehicle, total_ttc):
        self._shipping_nbr = shipping_nbr
        self._client_name = client_name
        self._shipping_address = shipping_address
        self._phone_nbr = phone_nbr
        self._order_nbr = order_nbr
        self._order_date = order_date
        self._central_dep_hr = central_dep_hr
        self._central_arr_hr = central_arr_hr
        self._worksite_arr_hr = worksite_arr_hr # retirer les heures d'arriver et de départ du lieu de livraison car il faut simplement le temps total qu'il a fallu mobiliser un vehicule
        self._worksite_dep_hr = worksite_dep_hr #
        self._product_id = product_id
        self._quantity = quantity
        self._product_type = product_type # à retirer car accessible depuis product_id
        self._vehicle = vehicle
        self._total_ttc = total_ttc # à retirer plus facil à calculer que de stocker
        
    def shipping_nbr(self):
        return self._shipping_nbr
    
    def shipping_nbr(self, shipping_nbr):
        self._shipping_nbr = shipping_nbr
        
    def client_name(self):
        return self._client_name
    
    def client_name(self, client_name):
        self._client_name = client_name
        
    def cin(self):
        return self._cin
    
    def cin(self, cin):
        self._cin = cin
        
    def shipping_address(self):
        return self._shipping_address
    
    def shipping_address(self, shipping_address):
        self._shipping_address = shipping_address
        
    def phone_nbr(self):
        return self._phone_nbr
    
    def phone_nbr(self, phone_nbr):
        self._phone_nbr = phone_nbr
        
    def order_nbr(self):
        return self._order_nbr
    
    def order_nbr(self, order_nbr):
        self._order_nbr = order_nbr
        
    def order_date(self):
        return self._order_date
    
    def order_date(self, order_date):
        self._order_date = order_date
        
    def central_dep_hr(self):
        return self._central_dep_hr
    
    def central_dep_hr(self, central_dep_hr):
        self._central_dep_hr = central_dep_hr
        
    def central_arr_hr(self):
        return self._central_arr_hr
    
    def central_arr_hr(self, central_arr_hr):
        self._central_arr_hr = central_arr_hr
        
    def worksite_arr_hr(self):
        return self._worksite_arr_hr
    
    def worksite_arr_hr(self, worksite_arr_hr):
        self._worksite_arr_hr = worksite_arr_hr
        
    def worksite_dep_hr(self):
        return self._worksite_dep_hr
    
    def worksite_dep_hr(self, worksite_dep_hr):
        self._worksite_dep_hr = worksite_dep_hr

    def product_id(self):
        return self._product_id

    def product_id(self, product_id):
        self._product_id = product_id
    def quantity(self):
        return self._quantity
    
    def quantity(self, quantity):
        self._quantity = quantity
        
    def product_type(self):
        return self._product_type
    
    def product_type(self, product_type):
        self._product_type = product_type

    def vehicle(self):
        return self._vehicle
    
    def vehicle(self, vehicle):
        self._vehicle = vehicle
        
    def total_ttc(self):
        return self._total_ttc
    
    def total_ttc(self, total_ttc):
        self._total_ttc = total_ttc

    def afficher_details(self):
        print(f"Shipping Number : {self._shipping_nbr}")
        print(f"Client Name : {self._client_name}")
        print(f"CIN : {self._cin}")
        print(f"Shipping Address : {self._shipping_address}")
        print(f"Phone Number : {self._phone_nbr}")
        print(f"Order Number : {self._order_nbr}")
        print(f"Order Date : {self._order_date}")
        print(f"Central Departure Hour : {self._central_dep_hr}")
        print(f"Central Arrival Hour : {self._central_arr_hr}")
        print(f"Worksite Arrival Hour : {self._worksite_arr_hr}")
        print(f"Worksite Departure Hour : {self._worksite_dep_hr}")
        print(f"Product ID : {self._product_id}")
        print(f"Quantity : {self._quantity}")
        print(f"Product Type : {self._product_type}")
        print(f"Vehicle : {self._vehicle}")
        print(f"Total TTC : {self._total_ttc}")

#general functions to get objects by id    
def getClientById(id):
    for client in client_instances:
        if client.clt_id==id:
            return client
    return None

def getOrderById(id):
    for order in order_instances:
        if order.order_id==id:
            return order
    return None

def getProductById(id):
    for prod in product_instances:
        if prod.id==id:
            return prod
    return None



###### TESTS 

# For Personne class
df1 = pd.read_csv("./test_class_personne.csv")
personne_instances = []
for index, row in df1.iterrows():
    personne_id = row['personne_id']
    name = row['name']
    address = row['address']
    phone_number = row['phone_number']
    personne = Personne(personne_id, name, address, phone_number)
    personne_instances.append(personne)
#for personne in personne_instances:
    #print(personne.id, personne.name, personne.address, personne.phone)

#For Supplier class
df2=pd.read_csv("./test_class_supplier.csv")
supplier_instances = []
for index, row in df2.iterrows():
    supplier_id = row['supplier_id']
    name = row['name']
    address = row['address']
    phone_number = row['phone_number']
    supplier = Supplier(supplier_id, name, address, phone_number)
    supplier_instances.append(supplier)
#for  supp in supplier_instances:
    #print(supp.id,supp.name,supp.address,supp.phone)

#For Employee class 
df3=pd.read_csv("./test_class_employee.csv")
employee_instances = []
for index, row in df3.iterrows():
    employee_id = row['employee_id']
    name = row['name']
    address = row['address']
    phone_number = row['phone_number']
    position=row['position']
    salary=row['salary']
    employee = Employee(employee_id, name, address, phone_number,position,salary)
    employee_instances.append(employee)
#for employee in employee_instances:
    #print(employee.id,employee.name,employee.address,employee.phone,employee.position,employee.salary)

#For Client class

df4=pd.read_csv("./test_class_client.csv")
client_instances = []
for index, row in df4.iterrows():
    client_id = row['client_id']
    name = row['name']
    address = row['address']
    phone_number = row['phone_number']
    client = Client(client_id, name, address, phone_number)
    client_instances.append(client)
#for client in client_instances:
    #print(client.id,client.name,client.address,client.phone)

#For product class
df5=pd.read_csv("./test_class_product.csv")
product_instances = []

for index, row in df5.iterrows():
    product_id = row['product_id']
    description = row['description']
    price = row['price']
    quantity_in_stock = row['quantity_in_stock']
    product = Product(product_id, description, price, quantity_in_stock)
    product_instances.append(product)

#for product in product_instances:
    #print(product.id, product.description, product.prix, product.stock)

#For Order class
df6=pd.read_csv("./test_class_order.csv")
order_instances = []

for index, row in df6.iterrows():
    order_id = row['order_id']
    order_date = row['order_date']
    client = row['client_id']
    products = row['product_ids']
    order = Order(order_id, order_date, client, products)
    order_instances.append(order)

#for order in order_instances:
    #print(order.id, order.date, order.client, order.product)

#For Invoice class

df7=pd.read_csv("./test_class_invoice.csv")
invoice_instances = []

for index, row in df7.iterrows():
    invoice_id = row['invoice_id']
    invoice_date = row['invoice_date']
    total_amount = row['total_amount']
    products = row['product_ids']
    invoice = Invoice(invoice_id, invoice_date, total_amount, products)
    invoice_instances.append(invoice)

#for invoice in invoice_instances:
    #print(invoice.id,invoice.date , invoice.total_amount, invoice.products)

#For Warehouse class
df8=pd.read_csv("./test_class_warehouse.csv")
warehouse=Warehouse(10000,0) #on a choisit pour le moment capacity=10000
prd=0
for index, row in df8.iterrows():
    quant = row['quantity_in_stock']
    prd+=quant

#For Sale class
df9 = pd.read_csv("test_class_sales.csv")
#print(df9)
sale_instances = []

for index, row in df9.iterrows():
    #sale_id = row['order_id']
    sale_date = row['sale_date']
    product_id = row['product_id']
    quantity_sold = row['quantity_sold']
    sale_price = row['sale_price']
    client_id = row['client_id']
    sale = Sale(sale_date, product_id, quantity_sold, sale_price,client_id)
    sale_instances.append(sale)



#### USER INTERFACE 

##########################

# Sample data (you can load data from CSV as well)
clients_data = [
    {"ID": 1, "Name": "Alice", "Email": "alice@example.com", "Phone Number": "1234567890"},
    {"ID": 2, "Name": "Bob", "Email": "bob@example.com", "Phone Number": "9876543210"},
    {"ID": 3, "Name": "Charlie", "Email": "charlie@example.com", "Phone Number": "5678901234"},
    {"ID": 4, "Name": "David", "Email": "david@example.com", "Phone Number": "8765432109"},
    {"ID": 5, "Name": "Eve", "Email": "eve@example.com", "Phone Number": "4321098765"}
]

orders_data = [
    {"Order ID": 1, "Order Date": "2023-07-18", "Client ID": 1, "Products": "Product A, Product B", "Type de Transaction": "CHEQUES", "Statut": "PAYE"},
    {"Order ID": 2, "Order Date": "2023-07-19", "Client ID": 2, "Products": "Product B, Product C", "Type de Transaction": "ESPECES", "Statut": "Non PAYE"},
    {"Order ID": 3, "Order Date": "2023-07-20", "Client ID": 3, "Products": "Product A, Product C", "Type de Transaction": "VIREMENT", "Statut": "AVANCE"},
    {"Order ID": 4, "Order Date": "2023-07-21", "Client ID": 4, "Products": "Product D, Product E", "Type de Transaction": "CHEQUES", "Statut": "PAYE"},
]

products_data = [
    {"Produit ID": 1, "Description": "test1", "Price": 15, "Quantité": 18, "Historique": "test1"},
    {"Produit ID": 2, "Description": "test2", "Price": 22, "Quantité": 4, "Historique": "test2"},
    {"Produit ID": 3, "Description": "test3", "Price": 37, "Quantité": 50, "Historique": "test3"},
    {"Produit ID": 4, "Description": "test4", "Price": 40, "Quantité": 45, "Historique": "test4"},
]
sales_data = [
    {"Sale ID": 10001,"Sale Date": date(2023,5,1), "Product ID": 1, "Quantity Sold": 5,"Sale Price":10.99,"Client ID":1001},
    {"Sale ID": 10002,"Sale Date": date(2023,5,2), "Product ID": 2, "Quantity Sold": 3,"Sale Price":15.99,"Client ID":1002},
    {"Sale ID": 10003,"Sale Date": date(2023,5,3), "Product ID": 3, "Quantity Sold": 10,"Sale Price":5.99,"Client ID":1003},
    {"Sale ID": 10004,"Sale Date": date(2023,5,4), "Product ID": 4, "Quantity Sold": 7,"Sale Price":8.99,"Client ID":1004},
    {"Sale ID": 10005,"Sale Date": date(2023,5,5), "Product ID": 5, "Quantity Sold": 2,"Sale Price":12.99,"Client ID":1005},
    {"Sale ID": 10006,"Sale Date": date(2023,5,6), "Product ID": 6, "Quantity Sold": 6,"Sale Price":7.99,"Client ID":1006},
    {"Sale ID": 10007,"Sale Date": date(2023,5,7), "Product ID": 7, "Quantity Sold": 9,"Sale Price":9.99,"Client ID":1007},
    {"Sale ID": 10008,"Sale Date": date(2023,5,8), "Product ID": 8, "Quantity Sold": 4,"Sale Price":6.99,"Client ID":1008},
    {"Sale ID": 10009,"Sale Date": date(2023,5,9), "Product ID": 9, "Quantity Sold": 1,"Sale Price":11.99,"Client ID":1009},
    {"Sale ID": 10010,"Sale Date": date(2023,5,10), "Product ID": 10, "Quantity Sold": 8,"Sale Price":14.99,"Client ID":1010}    
    ]  
sales_data_displayed = sales_data

# Modification et lecture des fichiers csv

def csv_list(chemin):
    df_livraison=pd.read_csv(chemin, sep = ';', encoding = 'latin-1')
    list_livraison = df_livraison.values.tolist()
    for i in range(len(list_livraison)):
        list_livraison[i] = list_livraison[i][0].split(';')
    return list_livraison

def add_livraison_csv(new_line, chemin):
    with open(chemin, 'a', newline='') as file:
        csv_writer = csv.writer(file)
        csv_writer.writerow(new_line)
        
   

livraison_data = csv_list("./test_livraison.csv")




# Fonction pour se connecter
def login(event=None): #event=None to be able to login by pressing the Return key or using the button
    employee_name = username_entry.get()
    employee_password = password_entry.get()
    if employee_name == ADMIN_USERNAME:
        if employee_password == ADMIN_PASSWORD:
            login_window.destroy()
            create_main_window()
        else:
            # Show an error message for mot de passe incorrect
            messagebox.showerror("Erreur", "Mot de passe incorrect")
    else:
        # Show an error message for nom d'utilisateur incorrect
        messagebox.showerror("Erreur", "Nom d'utilisateur incorrect")        


def validate_id(action, index, value_if_allowed, prior_value, text, validation_type, trigger_type, widget_name):
    """Validate the ID field to allow only numeric input."""
    if text.isdigit() or text == "":
        return True
    else:
        return False



def get_next_order_id():
    existing_ids = [order.order_id for order in order_instances]
    new_id = 1
    while new_id in existing_ids:
        new_id += 1
    return new_id

def refresh_order_ids():
    for index, order in enumerate(orders_data):
        order["Order ID"] = index + 1

def getOrderById(id):
    for order in order_instances:
        if order.order_id==id:
            return order
    messagebox.showerror("Erreur", "ID Commande Invalide")
    return None
        

def is_numeric_input(input_str):
    """Check if the input string is numeric."""
    return re.match(r'^\d+$', input_str) is not None

def validate_phone_number(action, index, value_if_allowed, prior_value, text, validation_type, trigger_type, widget_name):
    """Validate the phone number field to allow only numeric input."""
    if text.isdigit() or text == "":
        return True
    else:
        return False
    
def get_next_client_id():
    existing_ids = [client.clt_id for client in client_instances]
    new_id = 1
    while new_id in existing_ids:
        new_id += 1
    return new_id    

def refresh_client_ids():
    for index, client in enumerate(clients_data):
        client["ID"] = index + 1

def refresh_product_ids():
    for index, product in enumerate(products_data):
        product["ID"] = index + 1
        
def is_valid_phone_number(phone_number):
    # Validate the phone number using a regular expression
    phone_pattern = re.compile(r'^216\d{1,9}$')
    return phone_pattern.match(phone_number)

def is_valid_email(email):
    # Validate the email address using a regular expression
    email_pattern = re.compile(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$')
    return email_pattern.match(email)


def generate_new_supplier_id():
    # Retrieve existing supplier IDs from the CSV file
    csv_file_path = "./test_class_supplier.csv"  # Replace with the actual path to your CSV file
    existing_ids = set()
# Création de la fenêtre principale
def create_main_window():
    window = tk.Tk()
    window.title("Gestion d'entreprise")
    window.geometry("800x600+0+0")
      
    navbar = tk.Frame(window)
    navbar.pack()

    frame = tk.Frame(window)
    frame.pack()
    

        
    def Products():     
        Clear_widgets(frame)
        def display_products():
            if not products_tree.get_children():
                for product in products_data:
                    products_tree.insert("", tk.END, values=(product["Produit ID"], product["Description"], product["Price"], product["Quantité"], product["Historique"]))

        def refresh_products_table():
            products_tree.delete(*products_tree.get_children())
            for product in products_data:
                products_tree.insert("", tk.END, values=(product["Produit ID"], product["Description"], product["Price"], product["Quantité"], product["Historique"]))
        
        def double_click_product(event):
            selected_item = products_tree.focus()
            if selected_item:
                values = products_tree.item(selected_item, "values")
                if values:
                    product_id_entry.delete(0, tk.END)
                    description_entry.delete(0, tk.END)
                    price_entry.delete(0, tk.END)
                    quantity_entry.delete(0, tk.END)
                    historique_entry.delete(0, tk.END)
                    product_id_entry.insert(tk.END, values[0])
                    description_entry.insert(tk.END, values[1])
                    price_entry.insert(tk.END, values[2])
                    quantity_entry.insert(tk.END, values[3])
                    historique_entry.insert(tk.END, values[4])
                    
        def add_product():
            produit_id = product_id_entry.get()
            description = description_entry.get()
            price = price_entry.get()
            quantity_in_stock = quantity_entry.get()
            historique = historique_entry.get()
            if produit_id and description and price and quantity_in_stock and historique:
                if is_numeric_input(produit_id):
                    new_id = get_next_order_id()
                    
                    df = pd.read_csv('./test_class_product.csv')
                    size = df.shape[0] + 1
                    df.loc[size] = [new_id, description, price, quantity_in_stock, historique]
                    df.to_csv('./test_class_product.csv', index=False)
                
                    new_product = Product(new_id, description, price, quantity_in_stock, historique)
                    product_instances.append(new_product)
                    products_tree.insert("", tk.END, values=(new_id, description, price, quantity_in_stock, historique))
                    product_id_entry.delete(0, tk.END)
                    description_entry.delete(0, tk.END)
                    price_entry.delete(0, tk.END)
                    quantity_in_stock.delete(0, tk.END)  # Set default value for Type de Transaction
                else:
                    messagebox.showerror("Erreur", "L'identifiant de client doit être une valeur numérique.")
            else:
                messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
                
        def delete_product():
            selected_item = products_tree.selection()
            if selected_item:
                item_id = products_tree.item(selected_item)["values"][0]
                for product in products_data:
                    if product["Produit ID"] == item_id:
                        
                        
                        df = pd.read_csv('./test_class_product.csv', sep = ',')
                        df = df[df['product_id'] != item_id]
                        df.to_csv('./test_class_product.csv', index=False)
                        
                        
                        products_data.remove(product)
                        break
                products_tree.delete(selected_item)
                refresh_product_ids()
                refresh_products_table()
            else:
                messagebox.showwarning("Avertissement", "Veuillez sÃ©lectionner un produit Ã  supprimer.")
                
        def modify_product():
            selected_item = products_tree.selection()
            if selected_item:
                product_id = product_id_entry.get()
                description = description_entry.get()
                price = price_entry.get()
                quantity = quantity_entry.get()
                historique = historique_entry.get()

                if product_id and description and price and quantity and historique :
                    if is_numeric_input(product_id) and is_numeric_input(price) and is_numeric_input(quantity):
                        product_id = int(product_id)
                        selected_produit_id = products_tree.item(selected_item)["values"][0]
                        if selected_produit_id == product_id:
                            products_tree.item(selected_item, values=(product_id, description, price, quantity, historique))
                            
                            df = pd.read_csv('./test_class_product.csv')
                            colonne_index = 'product_id'
                            df = df.set_index(colonne_index)
                            
                            nouvelles_valeurs = {'product_id': product_id, 'description': description, 'price' : price, 'quantity_in_stock' : quantity, 'historique' : historique}
                            df.loc[product_id] = nouvelles_valeurs
                            df.reset_index(inplace = True)
                            
                            df.to_csv('./test_class_product.csv', index = False)
                            
                            messagebox.showinfo("Succès", "Commande modifiée avec succès.")
                        else:
                            messagebox.showerror("Erreur", "L'ID de la commande ne peut pas être modifié.")
                    else:
                        messagebox.showerror("Erreur", "L'identifiant de commande et l'identifiant de client doivent être des valeurs numériques.")
                else:
                    messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
            else:
                messagebox.showwarning("Avertissement", "Veuillez sélectionner une commande à modifier.")
    # Create the table to display products data
        titre_label = tk.Label(frame, text="Products", font=("Arial", 16))
        titre_label.pack(pady=5)
        
        columns_products = ("Produit ID", "Description", "Price", "Quantité", "Historique")
        
        tree_frame = tk.Frame(frame)
        tree_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical")
        scrollbar.pack(side='right', fill='y')
        
        style = ttk.Style()
        style.configure('Treeview', rowheight=25)
        products_tree = ttk.Treeview(tree_frame, columns=columns_products, show="headings", style='Custom.Treeview')
        
        scrollbar.config(command=products_tree.yview)  
        
        for col in columns_products:
            products_tree.heading(col, text=col)
            products_tree.column(col, width=5)

        products_tree.pack(fill=tk.BOTH, expand=True, pady=5)
        products_tree.bind("<Double-1>", double_click_product)
        
        # Add the labels and input fields for adding/modifying an order
        input_frame_products = tk.Frame(frame)
        input_frame_products.pack()

        product_id_label = tk.Label(input_frame_products, text="ID du produit :")
        product_id_label.pack(side=tk.LEFT, padx=5)
        product_id_entry = tk.Entry(input_frame_products)
        product_id_entry.pack(side=tk.LEFT, padx=5)

        description_label = tk.Label(input_frame_products, text="Description :")
        description_label.pack(side=tk.LEFT, padx=5)
        description_entry = tk.Entry(input_frame_products)
        description_entry.pack(side=tk.LEFT, padx=5)


        price_label = tk.Label(input_frame_products, text="Prix :")
        price_label.pack(side=tk.LEFT, padx=5)
        price_entry = tk.Entry(input_frame_products)
        price_entry.pack(side=tk.LEFT, padx=5)

        quantity_label = tk.Label(input_frame_products, text="Quantité :")
        quantity_label.pack(side=tk.LEFT, padx=5)
        quantity_entry = tk.Entry(input_frame_products)
        quantity_entry.pack(side=tk.LEFT, padx=5)

        historique_label = tk.Label(input_frame_products, text="Historique :")
        historique_label.pack(side=tk.LEFT, padx=5)
        historique_entry = tk.Entry(input_frame_products)
        historique_entry.pack(side=tk.LEFT, padx=5)
        
        button_frame_products = tk.Frame(frame)
        button_frame_products.pack(pady=5)
        
        display_products = ctk.CTkButton(button_frame_products, text="Liste des produits", command=display_products)
        display_products.pack(side=tk.LEFT, padx=5)
        # Add the buttons for adding/modifying a product


        add_button_products = tk.Button(button_frame_products, text="Ajouter Produit", command=add_product)
        add_button_products.pack(side=tk.LEFT, padx=5)

        delete_button_orders = tk.Button(button_frame_products, text="Supprimer Produit", command=delete_product)
        delete_button_orders.pack(side=tk.LEFT, padx=5)

        modify_button_orders = tk.Button(button_frame_products, text="Modifier Produit", command=modify_product)
        modify_button_orders.pack(side=tk.LEFT, padx=5)

        
    def Supplier():
        Clear_widgets(frame)
        def refresh_suppliers():
            # Clear old data from the display
            suppliers_tree.delete(*suppliers_tree.get_children())

            # Retrieve supplier data from the CSV file
            csv_file_path = "./test_class_supplier.csv"  

            with open(csv_file_path, newline='', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile)
                suppliers_data = list(reader)

            for supplier in suppliers_data:
                suppliers_tree.insert("", tk.END, values=(
                    supplier["supplier_id"], supplier["name"], supplier["address"],
                    supplier["phone_number"], supplier["contact_person"], supplier.get("email", "")
                ))

            # Clear the entry fields
            for entry in entries:
                entry.delete(0, tk.END)
        
        def double_click_supplier(event):
            selected_item = suppliers_tree.focus()
            if selected_item:
                values = suppliers_tree.item(selected_item, "values")
                if values:
                    # Clear the existing data from the specific entry variables
                    supplier_id_entry_supplier.delete(0, tk.END)
                    name_entry_supplier.delete(0, tk.END)
                    address_entry_supplier.delete(0, tk.END)
                    phone_entry_supplier.delete(0, tk.END)
                    contact_person_entry_supplier.delete(0, tk.END)
                    email_entry_supplier.delete(0, tk.END)

                    # Insert the values into the specific entry variables
                    supplier_id_entry_supplier.insert(tk.END, values[0])
                    name_entry_supplier.insert(tk.END, values[1])
                    address_entry_supplier.insert(tk.END, values[2])
                    phone_entry_supplier.insert(tk.END, values[3])
                    contact_person_entry_supplier.insert(tk.END, values[4])
                    email_entry_supplier.insert(tk.END, values[5])
        
        def add_supplier():
            try:
                global name_entry_supplier, address_entry_supplier, phone_entry_supplier, contact_person_entry_supplier, email_entry_supplier

                # Retrieve values from the entry fields using specific entry variables
                new_name = name_entry_supplier.get()
                new_address = address_entry_supplier.get()
                new_main_number = phone_entry_supplier.get()
                new_contact_person = contact_person_entry_supplier.get()
                new_email = email_entry_supplier.get()

                # Validate the phone number
                if not is_valid_phone_number(new_main_number):
                    messagebox.showerror("Erreur", "Le numéro de téléphone doit commencer par 216 et contenir de 1 à 9 chiffres.")
                    return  # Exit the function if the phone number is not valid

                # Validate the email address
                if not is_valid_email(new_email):
                    messagebox.showerror("Erreur", "L'adresse e-mail n'est pas valide.")
                    return  # Exit the function if the email address is not valid

                # Generate a new ID automatically
                new_id = generate_new_supplier_id()

                # Create a new supplier object
                new_supplier = Supplier(new_id, new_name, new_address, new_main_number, new_contact_person, new_email)

                # Insert the new supplier data into the CSV file
                csv_file_path = "./test_class_supplier.csv"

                with open(csv_file_path, mode='a', newline='', encoding='utf-8') as csv_file:
                    csv_writer = csv.writer(csv_file, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
                    csv_writer.writerow([
                        new_supplier._personne_id,  # Assuming this is the attribute for ID
                        new_supplier._name,
                        new_supplier._address,
                        new_supplier._phone_number,  # Assuming this is the attribute for phone number
                        new_supplier.contact_person,
                        new_supplier.email
                    ])

                # Update the display
                refresh_suppliers()
                messagebox.showinfo("Succès", "Fournisseur ajouté avec succès.")

            except Exception as e:
                messagebox.showerror("Erreur", f"Une erreur s'est produite : {str(e)}")
                
                
        def delete_supplier():
            # Retrieve the supplier ID to delete
            supplier_id_to_delete = int(supplier_id_entry_supplier.get())  # Assume the ID to delete is entered in the supplier ID field

            # Read the existing data from the CSV file
            csv_file_path = "./test_class_supplier.csv"

            with open(csv_file_path, mode='r', newline='', encoding='utf-8') as csv_file:
                csv_reader = csv.reader(csv_file)
                rows = list(csv_reader)

            deleted_supplier = None

            # Find and delete the supplier in the CSV data
            for idx, row in enumerate(rows):
                if idx > 0:  # Skip the header row
                    current_supplier_id = int(row[0])
                    if current_supplier_id == supplier_id_to_delete:
                        deleted_supplier = rows.pop(idx)
                        break

            if deleted_supplier:
                # Update the CSV file with the modified data
                with open(csv_file_path, mode='w', newline='', encoding='utf-8') as csv_file:
                    csv_writer = csv.writer(csv_file, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
                    csv_writer.writerows([rows[0]] + rows[1:])  # Write the header row and updated data

                # Update the display
                refresh_suppliers()

                # Clear the entry fields for the deleted supplier from the interface using specific entry variables
                supplier_id_entry_supplier.delete(0, tk.END)
                name_entry_supplier.delete(0, tk.END)
                address_entry_supplier.delete(0, tk.END)
                phone_entry_supplier.delete(0, tk.END)
                contact_person_entry_supplier.delete(0, tk.END)
                email_entry_supplier.delete(0, tk.END)

                # Inform about the deletion
                messagebox.showinfo("Suppression", f"Fournisseur {deleted_supplier[1]} supprimé avec succès.")
            else:
                messagebox.showerror("Erreur", f"Aucun fournisseur trouvé avec l'ID {supplier_id_to_delete}.")
                
                
        def modify_supplier():
            # Get the ID of the supplier to modify
            id_to_modify = int(supplier_id_entry_supplier.get())

            # Read the existing data from the CSV file
            csv_file_path = "./test_class_supplier.csv"

            with open(csv_file_path, mode='r', newline='', encoding='utf-8') as csv_file:
                csv_reader = csv.reader(csv_file)
                rows = list(csv_reader)

            existing_supplier = None
            idx_to_modify = None

            # Find the supplier in the CSV data
            for idx, row in enumerate(rows):
                if idx > 0:  # Skip the header row
                    current_supplier_id = int(row[0])
                    if current_supplier_id == id_to_modify:
                        existing_supplier = row
                        idx_to_modify = idx
                        break

            if existing_supplier is not None:
                # Get the updated values from the interface
                new_name = name_entry_supplier.get()
                new_address = address_entry_supplier.get()
                new_contact_person = contact_person_entry_supplier.get()
                new_email = email_entry_supplier.get()
                new_phone_number = phone_entry_supplier.get()

                # Validate the updated phone number if it's not empty
                if new_phone_number and not is_valid_phone_number(new_phone_number):
                    messagebox.showerror("Erreur", "Le numéro de téléphone doit commencer par 216 et contenir de 1 à 9 chiffres.")
                    return  # Exit the function if the phone number is not valid

                # Validate the updated email if it's not empty
                if new_email and not is_valid_email(new_email):
                    messagebox.showerror("Erreur", "L'adresse e-mail n'est pas valide.")
                    return  # Exit the function if the email address is not valid

                # Update the supplier details in the CSV data
                existing_supplier[1] = new_name
                existing_supplier[2] = new_address
                existing_supplier[4] = new_contact_person

                # Update phone number only if it's not empty and valid
                if new_phone_number:
                    existing_supplier[3] = new_phone_number

                # Update email only if it's not empty and valid
                if new_email:
                    existing_supplier[5] = new_email

                # Update the CSV file with the modified data
                with open(csv_file_path, mode='w', newline='', encoding='utf-8') as csv_file:
                    csv_writer = csv.writer(csv_file, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
                    csv_writer.writerows([rows[0]] + rows[1:])  # Write the header row and updated data

                # Refresh the display
                refresh_suppliers()
                messagebox.showinfo("Succès", "Détails du fournisseur mis à jour avec succès.")
            else:
                messagebox.showerror("Erreur", "Aucun fournisseur trouvé avec l'ID spécifié.")


        # Create the table to display suppliers data
        titre_label = tk.Label(frame, text="Fournisseurs", font=("Arial", 16))
        titre_label.pack(pady=5)

        columns_suppliers = ("Supplier ID", "Name", "Address", "Phone Number", "Contact Person", "Email")

        tree_frame_suppliers = tk.Frame(frame)
        tree_frame_suppliers.pack(fill=tk.BOTH, expand=True, pady=10)

        scrollbar_suppliers = ttk.Scrollbar(tree_frame_suppliers, orient="vertical")
        scrollbar_suppliers.pack(side='right', fill='y')

        style_suppliers = ttk.Style()
        style_suppliers.configure('Treeview', rowheight=25)

        suppliers_tree = ttk.Treeview(tree_frame_suppliers, columns=columns_suppliers, show="headings", style='Custom.Treeview')

        scrollbar_suppliers.config(command=suppliers_tree.yview)

        for col in columns_suppliers:
            suppliers_tree.heading(col, text=col)
            suppliers_tree.column(col, width=150)

        suppliers_tree.pack(fill=tk.BOTH, expand=True, pady=10)
        suppliers_tree.bind("<Double-1>", double_click_supplier)

        # Add the labels and input fields for adding/modifying a supplier
        input_frame_suppliers = tk.Frame(frame)
        input_frame_suppliers.pack()

        fields = ["ID du Fournisseur", "Nom du Fournisseur", "Adresse du Fournisseur",
                "Numéro de téléphone du Fournisseur", "Nom du Contact", "Email du Fournisseur"]

        entries = []

        for idx, field_text in enumerate(fields):
            row_num = idx // 2
            col_num = idx % 2 * 2 + 1

            label = tk.Label(input_frame_suppliers, text=f"{field_text} :")
            label.grid(row=row_num, column=col_num - 1, padx=5, pady=5, sticky="e")

            entry = tk.Entry(input_frame_suppliers)
            entry.grid(row=row_num, column=col_num, padx=5, pady=5)
            entries.append(entry)

        # Access the entry fields using the list entries with specific names
        supplier_id_entry_supplier = entries[0]
        name_entry_supplier = entries[1]
        address_entry_supplier = entries[2]
        phone_entry_supplier = entries[3]
        contact_person_entry_supplier = entries[4]
        email_entry_supplier = entries[5]


        # Add the buttons for adding/modifying a supplier
        button_frame_suppliers = tk.Frame(frame)
        button_frame_suppliers.pack(pady=10)
        
        display_refresh_suppliers = tk.Button(button_frame_suppliers, text="Liste des fournisseurs", command=refresh_suppliers)
        display_refresh_suppliers.pack(side=tk.LEFT, padx=5)

        add_button_suppliers = tk.Button(button_frame_suppliers, text="Ajouter Fournisseur", command=add_supplier)
        add_button_suppliers.pack(side=tk.LEFT, padx=5)

        delete_button_suppliers = tk.Button(button_frame_suppliers, text="Supprimer Fournisseur", command=delete_supplier)
        delete_button_suppliers.pack(side=tk.LEFT, padx=5)

        modify_button_suppliers = tk.Button(button_frame_suppliers, text="Modifier Fournisseur", command=modify_supplier)
        modify_button_suppliers.pack(side=tk.LEFT, padx=5)



        button_refresh_suppliers = tk.Button(button_frame_suppliers, text="Rafraîchir", command=refresh_suppliers)
        button_refresh_suppliers.pack(side=tk.LEFT, padx=5)

    def Sales():
        Clear_widgets(frame)
        tk.Label(frame,text="Sales").pack()
        
    def Deliveries():
        Clear_widgets(frame)
        tk.Label(frame,text="Deliveries").pack()
        
        
    def Clients():
        Clear_widgets(frame)
        tk.Label(frame,text="Clients").pack()
        
        
    #//////////////////////////// INTERFACE ORDER ///////////////////////////////////////////
        
    def orders():    
        Clear_widgets(frame)

        def add_order():
            client_id = order_client_id_entry.get()
            if client_id.isnumeric():
                client = getClientById(int(client_id))
                if client:
                    order_date = order_date_label.cget('text')
                    type_transaction = type_transaction_var.get()
                    statut = statut_var.get()
                    pompe = True if pompe_var.get() == 'Oui' else False

                    if order_date != "" and order_listprod_var.get()!="":
                        products = [(getProductById(int(id)), qty) for id, qty in (a.split(' : ') for a in order_listprod_var.get().split('\n'))]
                        new_id = get_next_order_id()
                        new_order = Order(new_id, order_date, client, products, type_transaction, pompe=pompe)
                        order_instances.append(new_order)
                        orders_tree.insert("", tk.END, values=(new_id, order_date, client.id, " | ".join([f"{prod.id} : {qty}" for prod, qty in products]), type_transaction))
                        order_date_label.config(text="")
                        order_client_id_entry.delete(0, tk.END)
                        order_listprod_var.set("")
                        
                        type_transaction_var.set("CHEQUE")  # Set default value for Type de Transaction
                        statut_var.set("PAYE")  # Set default value for Statut

                        df = pd.read_csv('./test_class_order.csv')
                        size = df.shape[0] + 1
                        df.loc[size] = [new_id, order_date, client_id, products]
                        df.to_csv('./test_class_order.csv', index=False)
                    else:
                        messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
                else:
                    # Le client n'existe pas, afficher un message d'erreur
                    messagebox.showerror("Client inexistant", "Le client n'existe pas. Veuillez le créer avant d'ajouter la commande.")
            else:
                messagebox.showerror("Erreur", "L'identifiant de client doit être une valeur numérique.")

        def modify_order():
            selected_item = orders_tree.selection()
            if selected_item:
                order_id = order_id_entry.get()
                order_date = order_date_label.cget('text')
                client_id = order_client_id_entry.get()
                type_transaction = type_transaction_var.get()
                statut = statut_var.get()

                if order_id and order_date and client_id and order_listprod_var.get()!="" and type_transaction and statut:
                    if is_numeric_input(order_id) and is_numeric_input(client_id):
                        order_id = int(order_id)
                        products = [(getProductById(int(id)), qty) for id, qty in (a.split(' : ') for a in order_listprod_var.get().split('\n'))]
                        selected_order_id = orders_tree.item(selected_item)["values"][0]
                        if selected_order_id == order_id:
                            
                            df = pd.read_csv('./test_class_order.csv')
                            colonne_index = 'order_id'
                            df = df.set_index(colonne_index)
                            
                            nouvelles_valeurs = {'order_date': order_date, 'client_id': client_id, 'product_ids' : products}
                            df.loc[order_id] = nouvelles_valeurs
                            df.reset_index(inplace = True)
                            
                            df.to_csv('./test_class_order.csv', index = False)
                            
                            orders_tree.item(selected_item, values=(order_id, order_date, client_id, " | ".join([f"{prod.id} : {qty}" for prod, qty in products]), type_transaction, statut))
                            messagebox.showinfo("Succès", "Commande modifiée avec succès.")
                        else:
                            messagebox.showerror("Erreur", "L'ID de la commande ne peut pas être modifié.")
                    else:
                        messagebox.showerror("Erreur", "L'identifiant de commande et l'identifiant de client doivent être des valeurs numériques.")
                else:
                    messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
            else:
                messagebox.showwarning("Avertissement", "Veuillez sélectionner une commande à modifier.")
        
        def delete_order():
            selected_item = orders_tree.selection()
            if selected_item:
                order_id = orders_tree.item(selected_item)["values"][0]
                for order in orders_data:
                    if order["Order ID"] == order_id:
                        
                        
                        df = pd.read_csv('./test_class_order.csv', sep = ',')
                        df = df[df['order_id'] != order_id]
                        df.to_csv('./test_class_order.csv', index=False)
                        
                        orders_data.remove(order)
                        break
                orders_tree.delete(selected_item)
                refresh_order_ids()
            else:
                messagebox.showwarning("Avertissement", "Veuillez sélectionner une commande à supprimer.")
        
        def double_click_order(event):
            selected_item = orders_tree.focus()
            if selected_item:
                values = orders_tree.item(selected_item, "values")
                if values:
                    order_id_entry.delete(0, tk.END)
                    order_client_id_entry.delete(0, tk.END)
                    order_id_entry.insert(tk.END, values[0])
                    order_date_label.config(text=values[1])
                    order_client_id_entry.insert(tk.END, values[2])
                    order_listprod_var.set("\n".join(values[3].split(" | ")))
        
        def GetDate(event): #create a new window with a calendar
            newWindow = tk.Toplevel(frame)
            newWindow.geometry('250x200+300+100')
            today = date.today()
            cal = Calendar(newWindow, selectmode = 'day', year = today.year, month = today.month, day = today.day, date_pattern='dd/mm/y')
            cal.pack(pady=5)
            
            def onclick_date(event):
                order_date_label.config(text=cal.get_date())
                newWindow.destroy()
            
            for row in cal._calendar:
                for lbl in row:
                    lbl.bind('<Double-1>', onclick_date)
        
        def ProductList(event): #create a new window to headle the product/qty list
            def onclickProduct(event=None):
                if order_add_qty_entry.get():
                    s = order_listprod_var.get()+'\n' if order_listprod_var.get()!="" else ""
                    s+=f"{order_add_product.get().split(')')[0][1:]} : {order_add_qty_entry.get()}"
                    order_listprod_var.set(s)
                    order_add_qty_entry.delete(0, tk.END)
                    order_add_qty_entry.focus_set()
                
            newWindow = tk.Toplevel(frame)
            newWindow.geometry('250x150+400+200')
            tk.Label(newWindow, text='ID Produit : Quantité (kg/m\u00B3)').grid(column=0, row=0) #^3 : U+00B3
            tk.Label(newWindow, textvariable=order_listprod_var, bg='white', bd=1, justify='left', anchor='w', relief='sunken').grid(column=0, row=1, columnspan=2, sticky='WE', pady=3)
            order_add_product = tk.StringVar(newWindow, f"({product_instances[0].id}) {product_instances[0].description}")
            ttk.Combobox(newWindow, textvariable= order_add_product, values=[f"({prod.id}) {prod.description}" for prod in product_instances], state="readonly").grid(column=0, row=2)
            order_add_qty_entry = tk.Entry(newWindow, width=14, validate='key', validatecommand=(frame.register(validate_id), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
            order_add_qty_entry.grid(column=1, row=2)
            order_add_qty_entry.bind('<Return>', onclickProduct)
            tk.Button(newWindow, text="Clear list", command=lambda:order_listprod_var.set("")).grid(column=0, row=3)
            tk.Button(newWindow, text="Ajouter produit", command=onclickProduct).grid(column=1, row=3, pady=5)

        def display_orders():
            if not orders_tree.get_children():
                for order in orders_data:
                    orders_tree.insert("", tk.END, values=(
                        order["Order ID"], order["Order Date"], order["Client ID"], order["Products"],
                        order.get("Type de Transaction", ""), order.get("Statut", "")
                    ))
        
        titre_label = tk.Label(frame, text="Orders", font=("Arial", 16))
        titre_label.pack(pady=5)

        columns = ("Order ID", "Date de la commande", "Client ID", "Produits", "Type de Transaction", "Statut")
        orders_tree = ttk.Treeview(frame, columns=columns, show="headings")

        for col in columns:
            orders_tree.heading(col, text=col)
            orders_tree.column(col, width=150)

        orders_tree.pack(fill=tk.BOTH, expand=True, pady=10)
        orders_tree.bind("<Double-1>", double_click_order)
        
        
        # Add the labels and input fields for adding/modifying an order
        input_frame_orders = tk.Frame(frame)
        input_frame_orders.pack()

        tk.Label(input_frame_orders, text="ID de commande :").pack(side=tk.LEFT, padx=5)
        order_id_entry = tk.Entry(input_frame_orders, validate="key")
        order_id_entry.config(validatecommand=(frame.register(validate_id), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
        order_id_entry.pack(side=tk.LEFT, padx=5)

        tk.Label(input_frame_orders, text="Date de commande :").pack(side=tk.LEFT, padx=5)
        order_date_label = tk.Label(input_frame_orders, bg="white", width=10, relief='sunken', bd=1, cursor='hand2')
        order_date_label.pack(side=tk.LEFT, padx=5)
        order_date_label.bind('<Button-1>', GetDate)
        
        tk.Label(input_frame_orders, text="ID de client :").pack(side=tk.LEFT, padx=5)
        order_client_id_entry = tk.Entry(input_frame_orders, validate="key")
        order_client_id_entry.config(validatecommand=(frame.register(validate_id), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
        order_client_id_entry.pack(side=tk.LEFT, padx=5)
        
        tk.Label(input_frame_orders, text="Produits :").pack(side=tk.LEFT, padx=5)
        order_listprod_var = tk.StringVar(input_frame_orders, "")
        order_listprod_label = tk.Label(input_frame_orders, textvariable=order_listprod_var, bg="white", width=40, relief='sunken', bd=1, justify='left', anchor='w', cursor='hand2')
        order_listprod_label.pack(side=tk.LEFT, padx=5)
        order_listprod_label.bind('<Button-1>', ProductList)
        
        
        # Add the select boxes for Type de Transaction and Statut
        transaction_frame = tk.Frame(frame)
        transaction_frame.pack(pady=10)

        tk.Label(transaction_frame, text="Type de Transaction:").pack(side=tk.LEFT, padx=5)
        type_transaction_var = tk.StringVar(frame)
        type_transaction_var.set("CHEQUES")
        type_transaction_select = ttk.Combobox(transaction_frame, textvariable=type_transaction_var, values=["CHEQUES", "ESPECES", "VIREMENT"], state="readonly")
        type_transaction_select.pack(side=tk.LEFT, padx=5)

        tk.Label(transaction_frame, text="Statut:").pack(side=tk.LEFT, padx=5)
        statut_var = tk.StringVar(frame)
        statut_var.set("PAYE")
        statut_select = ttk.Combobox(transaction_frame, textvariable=statut_var, values=["PAYE", "Non PAYE", "AVANCE"], state="readonly")
        statut_select.pack(side=tk.LEFT, padx=5)

        tk.Label(transaction_frame, text="Pompé :").pack(side=tk.LEFT, padx=5)
        pompe_var = tk.StringVar(frame)
        pompe_var.set("Oui")
        ttk.Combobox(transaction_frame, textvariable=pompe_var, values=["Oui", "Non"], state="readonly").pack(side=tk.LEFT, padx=5)
                        
        # Add the buttons for adding/modifying an order
        button_frame_orders = tk.Frame(frame)
        button_frame_orders.pack(pady=10)
        
        #buttons
        ctk.CTkButton(button_frame_orders, text="Liste des commandes", command=display_orders).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame_orders, text="Ajouter Commande", command=add_order).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame_orders, text="Supprimer Commande", command=delete_order).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame_orders, text="Modifier Commande", command=modify_order).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame_orders, text="Créer un Bon de Commande", command=lambda: getOrderById(int(order_id_entry.get())).CreateBDC()).pack(side=tk.LEFT, padx=5)
        

        #////////////////////////////  FIN INTERFACE ORDER ///////////////////////////////////////////
              
        
        
    ctk.CTkButton(navbar, text="Orders", command=orders).pack(side="left")
    ctk.CTkButton(navbar, text="Clients", command=Clients).pack(side="left")
    ctk.CTkButton(navbar, text="Products", command=Products).pack(side="left")
    ctk.CTkButton(navbar, text="Sales", command=Sales).pack(side="left")
    ctk.CTkButton(navbar, text="Deliveries", command=Deliveries).pack(side="left")
    ctk.CTkButton(navbar, text="Suppliers", command=Supplier).pack(side="left")


    #open_section(frame) #I had an error with this line
    window.mainloop()

# Fonction pour ouvrir différentes sections
def open_section(section_name, frame):
    Clear_widgets(frame)
    tk.Label(frame, text=section_name).pack()

# Fonction pour vider le contenu d'un cadre
def Clear_widgets(frame):
    for widget in frame.winfo_children():
        widget.destroy()

# Paramètres d'administration
ADMIN_USERNAME = "admin"
ADMIN_PASSWORD = "password"

# Fenêtre de connexion

login_window = ctk.CTk()
login_window.title("Connexion")
login_window.geometry("600x600")
login_window.configure(bg="navy")  # Set the background color to navy blue

# Create a frame for the login elements
login_frame = ctk.CTkFrame(login_window)  
login_frame.pack(pady=(50, 20), padx=20)  # Add padx to space the frame from the window edge

# Ajouter l'image
image = Image.open("./téléchargement.png")
photo = ImageTk.PhotoImage(image)
image_label = tk.Label(login_frame, image=photo, bg="#FF6F61")  # Rouge pas foncé
image_label.grid(row=0, column=0, columnspan=2)

# Add labels and entry fields
username_label = ctk.CTkLabel(login_frame, text="Nom d'utilisateur :", font=("Helvetica", 12))  
username_label.grid(row=1, column=0, pady=(0, 5), sticky="w")

username_entry = ctk.CTkEntry(login_frame, font=("Helvetica", 12))
username_entry.grid(row=2, column=0, pady=(0, 10), sticky="w")  # Adjusted row and pady

password_label = ctk.CTkLabel(login_frame, text="Mot de passe :", font=("Helvetica", 12))  
password_label.grid(row=3, column=0, pady=(0, 5), sticky="w")

password_entry = ctk.CTkEntry(login_frame, show="*", font=("Helvetica", 12))
password_entry.grid(row=4, column=0, pady=(0, 10), sticky="w")
password_entry.bind('<Return>', login)

login_button = ctk.CTkButton(login_window, text="Se connecter", command=login, font=("Helvetica", 12))
login_button.pack(pady=(0, 20))

login_window.mainloop()