#imports
import pandas as pd
from PIL import Image, ImageTk
import tkinter as tk
from tkinter import ttk, simpledialog, messagebox
import re
import customtkinter as ctk 
from datetime import date, datetime, timedelta
import csv
import datetime
from tkcalendar import DateEntry, Calendar
import os

#   pip install python-docx (for editing words)
from docx import Document
from docx.shared import Pt, Inches, RGBColor #text size, color
from docx.enum.text import WD_ALIGN_PARAGRAPH #paragraph alignment
import comtypes.client

ctk.deactivate_automatic_dpi_awareness()
ctk.set_appearance_mode("Dark")  # Modes: "System" (standard), "Dark", "Light"
ctk.set_default_color_theme("blue")

class Personne:
    def __init__(self, personne_id, name=None, address=None, phone_number=None):
        self._personne_id = personne_id
        self._name = name
        self._address = address
        self._phone_number = phone_number
    
    def get_id(self):
        return self._personne_id
    
    def set_id(self, personne_id):
        self._personne_id = personne_id
    id=property(get_id,set_id)
    
    def get_name(self):
        return self._name
    def set_name(self, name):
        self._name = name
    name=property(get_name,set_name)
    
    def get_address(self):
        return self._address
    
    def set_address(self, address):
        self._address = address
    address=property(get_address,set_address)
    
    def get_phone_number(self):
        return self._phone_number
    
    def set_phone_number(self, phone_number):
        self._phone_number = phone_number
    phone=property(get_phone_number,set_phone_number)

class Supplier(Personne):
    
    def __init__(self, supplier_id, name=None, address=None, phone_number=None):
        super().__init__(supplier_id, name, address, phone_number)
        self._supplier_id = supplier_id
    
    def get_supplier_id(self):
        return self._supplier_id
    
    def set_supplier_id(self, supplier_id):
        self._supplier_id = supplier_id
    id=property(get_supplier_id,set_supplier_id)

class Employee(Personne):
    def __init__(self, employee_id, name=None, address=None, phone_number=None, position=None, salary=None, authority=None):
        super().__init__(employee_id, name, address, phone_number)
        self._employee_id = employee_id
        self._position = position
        self._salary = salary
        self.phone_number = phone_number
        self.name = name
        self._authority = authority

    def get_employee_id(self):
        return self._employee_id

    def set_employee_id(self, employee_id):
        self._employee_id = employee_id
    emp_id = property(get_employee_id, set_employee_id)

    def get_position(self):
        return self._position

    def set_position(self, position):
        self._position = position
    position = property(get_position, set_position)

    def get_salary(self):
        return self._salary

    def set_salary(self, salary):
        self._salary = salary
    salary = property(get_salary, set_salary)

    def get_phone_number(self):
        return self.phone_number

    def set_phone_number(self, phone):
        self.phone_number = phone
    phone = property(get_phone_number, set_phone_number)

    def get_authority(self):
        return self._authority

    def set_authority(self, authority):
        self._authority = authority
    authority = property(get_authority, set_authority)

class Client(Personne):
    def __init__(self, client_id, name=None, address=None, phone_number=None, email=None):
        super().__init__(client_id, name, address, phone_number)
        self._client_id = client_id
        self._email = email
    def get_client_id(self):
        return self._client_id
    def set_client_id(self, client_id):
        self._client_id = client_id

    clt_id=property(get_client_id,set_client_id)

    def get_email(self):
        return self._email
    def set_email(self, email):
        self._email = email
    email=property(get_email,set_email)

class Database:
    def __init__(self):
        self.employees = []
        self.products = []
        self.orders = []
        self.clients = []
        self.warehouse = Warehouse(capacity=1000) #on peut changer la capacité si nécessaire 
        self.sales = []
        self.suppliers = []
        self.invoices = []
        self.reports = []

    def add_employee(self, employee):
        self.employees.append(employee)

    def add_product(self, product):
        self.products.append(product)

    def add_order(self, order):
        self.orders.append(order)

    def add_client(self, client):
        self.clients.append(client)

    def add_sale(self, sale):
        self.sales.append(sale)

    def add_supplier(self, supplier):
        self.suppliers.append(supplier)

    def add_invoice(self, invoice):
        self.invoices.append(invoice)

    def add_report(self, report):
        self.reports.append(report)

    def show_database(self):
        print("Employees:")
        for employee in self.employees:
            print(employee.get_employee_id(), employee.get_name(), employee.get_position())
        
        print("\nProducts:")
        for product in self.products:
            print(product.get_product_id(), product.get_name(), product.get_price())
        
        print("\nOrders:")
        for order in self.orders:
            print(order.get_order_id(), order.get_client_id(), order.get_products(), order.get_order_date())
        
        print("\nClients:")
        for client in self.clients:
            print(client.get_client_id(), client.get_name(), client.get_address())
        
        print("\nSales:")
        for sale in self.sales:
            print(sale.get_sale_id(), sale.get_product_id(), sale.get_quantity(), sale.get_sale_date())
        
        print("\nSuppliers:")
        for supplier in self.suppliers:
            print(supplier.get_supplier_id(), supplier.get_name(), supplier.get_address())
        
        print("\nInvoices:")
        for invoice in self.invoices:
            print(invoice.get_invoice_id(), invoice.get_client_id(), invoice.get_invoice_date(), invoice.get_total_amount())
        
        print("\nReports:")
        for report in self.reports:
            print(report.get_report_id(), report.get_title(), report.get_report_date(), report.get_data())

class Product:
    def __init__(self, product_id, description=None, price=None, quantity_in_stock=None, historique = None):
        self._product_id = product_id
        self._description = description
        self._price = price
        self._quantity_in_stock = quantity_in_stock
        self.historique = historique
    
    def get_product_id(self):
        return self._product_id
    
    def set_product_id(self, product_id):
        self._product_id = product_id
    id=property(get_product_id,set_product_id)
    
    def get_description(self):
        return self._description
    
    def set_description(self, description):
        self._description = description
    description=property(get_description,set_description)

    def get_price(self):
        return self._price
    
    def set_price(self, price):
        self._price = price
    prix=property(get_price,set_price)
    
    def get_quantity_in_stock(self):
        return self._quantity_in_stock
    
    def set_quantity_in_stock(self, quantity_in_stock):
        self._quantity_in_stock = quantity_in_stock

    stock=property(get_quantity_in_stock,set_quantity_in_stock)
    
class Order:
    def __init__(self, order_id, order_date, client, products, payment_type="", price=0, price_paid = 0, pompe=False, statut="Undefined"):
        self._order_id = order_id
        self._order_date = order_date
        self._client = client
        self._products = products #list( (Product obj, qty), ... )
        self._price=round(sum(product[0].prix*int(product[1]) for product in self._products),2)
        self._price_paid = price_paid
        self._payment_type=payment_type.capitalize()
        self._pompe=pompe
        self._statut=statut
    
    def get_order_id(self):
        return self._order_id
    def set_order_id(self, order_id):
        self._order_id = order_id
    order_id=property(get_order_id,set_order_id)
    
    def get_order_date(self):
        return self._order_date
    def set_order_date(self, order_date):
        self._order_date = order_date
    order_date=property(get_order_date,set_order_date)
    
    def get_client(self):
        return self._client
    def set_client(self, client):
        self._client = client
    client=property(get_client,set_client)
    
    def get_products(self):
        return self._products
    def set_products(self, products):
        self._products = products
    products=property(get_products,set_products)
    
    def get_price(self):
        return self._price
    def set_price(self, price):
        self._price=price
    price=property(get_price, set_price)
    
    def get_price_paid(self):
        return self._price_paid
    def set_price_paid(self, price_paid):
        self._price_paid=price_paid
    price_paid=property(get_price_paid, set_price_paid)
    
    def get_payment_type(self):
        return self._payment_type
    def set_payment_type(self, payment_type):
        self._payment_type=payment_type.capitalize()
    payment_type=property(get_payment_type,set_payment_type)
    
    def get_pompe(self):
        return self._pompe
    def set_pompe(self, pompe):
        self._pompe=pompe
    pompe=property(get_pompe, set_pompe)
    
    def get_statut(self):
        return self._statut
    def set_statut(self, statut):
        self._statut=statut
    statut=property(get_statut, set_statut)
    
    def get_str_Products(self, id=False):
        if id:
            return " | ".join([f"{prod.id} : {qty}" for prod, qty in self._products])
        else:
            return " | ".join([f"{prod.description} : {qty}" for prod, qty in self._products])
    
    def CreateBDC(self):
        titleSize = Pt(13)
        valueSize = Pt(12)
            
        def AddTitleValue(paragraph, title, value="", length=0, alignment='left'):
            if alignment=='right':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment=='center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = paragraph.add_run(title)
            run.bold=True
            run.font.size=titleSize
            paragraph.add_run(str(value)).font.size = valueSize
            if len(str(value))<length and len(str(value))!=0:
                paragraph.add_run(" "*(length-len(str(value)))).font.size = valueSize
        
        document = Document('BON DE COMMANDE - blank.docx')
        
        #date
        AddTitleValue(document.paragraphs[-1], 'CM DU : ', date.today().strftime("%d/%m/%Y"), alignment='right')
        
        #id commande
        AddTitleValue(document.add_paragraph(), 'BON DE COMMANDE N° : ', self.order_id, 28)
        
        #nom client
        p = document.add_paragraph()
        AddTitleValue(p, 'Client : ', self.client.name, 27)
        
        #tel client
        AddTitleValue(p, '\t\t\tTél : ', self.client.phone, 24)
        
        #email client
        AddTitleValue(p, '\t\t\tEmail : ', self.client.email, 24)
        
        #products
        table = document.add_table(rows=1, cols=3)
        """
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Produit'
        hdr_cells[1].text = 'Quantité'
        hdr_cells[2].text = "Prix à l'unité"
        
        for product in self.products:
            row_cells = table.add_row().cells
            row_cells[0].text = product[0].description
            row_cells[1].text = str(product[1])
            row_cells[2].text = str(product[0].prix)
        """
        hdr_cells = table.rows[0].cells
        for cell, text in zip(hdr_cells, ['Produits', 'Quantité', "Prix à l'unité"]):
            cell.text = text
            # Set font properties (bold and size)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = titleSize  # Adjust size as needed

        # Populate the table with data from the list of products and set font size
        for product in self.products:
            row_cells = table.add_row().cells
            for cell, data in zip(row_cells, [product[0].description, str(product[1]), str(product[0].prix)]):
                cell.text = str(data)
                cell.paragraphs[0].runs[0].font.size = valueSize
        
        # dosage
        p = document.add_paragraph()
        AddTitleValue(p, 'Dosage :\t', 'Pompé' if self.pompe else 'Non pompé')
        
        #payment method
        if self.payment_type=='Especes':
            AddTitleValue(p, '\t\t\tRèglement : ','Espèces')
        elif self.payment_type=='Cheque':
            AddTitleValue(p, '\t\t\tRèglement : ','Chèque')
        else:
            AddTitleValue(p, '\t\t\tRèglement : ','Virement')

        #date commande
        p = document.add_paragraph()
        AddTitleValue(p, 'Date de livraison : ', self.order_date, 24)

        #adresse client
        AddTitleValue(p, '\t\t\t\tDestination : ', self.client.address, 28)
        
        #montant
        p = document.add_paragraph()
        AddTitleValue(p, 'Montant : ', self.price, 19)
        
        #reste
        AddTitleValue(p, '\t\t\t\t\t\tReste à payer : ', self.price-self.price_paid, 19)
        
        #signatures
        AddTitleValue(document.add_paragraph(), 'Signature du client :\t\t\t\t\t\t\t\t\tSignature du Chef de la Centrale : ', "")
        try:
            if not os.path.isdir('./Bon de Commandes'):
                os.mkdir('./Bon de Commandes')
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run().add_picture('signature.png')
        except Exception as e:
            print(e)
        document.save(f'./Bon de Commandes/Bon de Commande - {self.order_id}.docx')

class Warehouse:
    def __init__(self, capacity=None,quantity_in_stock=None,products_in_stock=None):
        self._capacity = capacity
        self._products_in_stock = products_in_stock
        self._quantity_in_stock=quantity_in_stock

    def get_capacity(self):
        return self._capacity
    
    def set_capacity(self, capacity):
        self._capacity = capacity
    capacity=property(get_capacity,set_capacity)
    
    def get_products_in_stock(self):
        return self._products_in_stock
    
    def add_product(self, product):
        self._products_in_stock.append(product)
    
    def remove_product(self, product):
        self._products_in_stock.remove(product)
    products=property(get_products_in_stock,add_product,remove_product)

    def get_quantity_in_stock(self):
        return self._quantity_in_stock
    def set_quantity_in_stock(self,q):
        self._quantity_in_stock=q
    quantity=property(get_quantity_in_stock,set_quantity_in_stock)

class Sale:
    def __init__(self,client,sale_id=None, sale_date=None, product=None, quantity_sold=None, sale_price=None):
        self._sale_date = sale_date
        self._product = product
        self._quantity_sold = quantity_sold
        self._sale_price = sale_price
        self._client = client

    def get_sale_id(self):
        return self._sale_id
    
    def set_sale_id(self, sale_id):
        self._sale_id = sale_id
    sale_id=property(get_sale_id,set_sale_id)
    
    def get_sale_date(self):
        return self._sale_date
    
    def set_sale_date(self, sale_date):
        self._sale_date = sale_date
    sale_date=property(get_sale_date,set_sale_date)
    
    def get_product(self):
        return self._product
    
    def set_product(self, product):
        self._product = product
    prod=property(get_product,set_product)
    
    def get_quantity_sold(self):
        return self._quantity_sold
    
    def set_quantity_sold(self, quantity_sold):
        self._quantity_sold = quantity_sold
    sold=property(get_quantity_sold,set_quantity_sold)
    
    def get_sale_price(self):
        return self._sale_price
    
    def set_sale_price(self, sale_price):
        self._sale_price = sale_price
    price=property(get_sale_price,set_sale_price)
    
    def get_client(self):
        return self._client
    
    def set_client(self, client):
        self._client = client
    client=property(get_client,set_client)

    def report(self, products):
        sales_per_client = {}
        for sale in self.sales:
            client = sale.get_client().get_client_id()
            quantity_sold = sale.get_quantity_sold()
            if client in sales_per_client:
                sales_per_client[client] += quantity_sold
            else:
                sales_per_client[client] = quantity_sold

        average_sales_per_client = {client: total_sales / len(self.sales) for client, total_sales in sales_per_client.items()}

        sales_per_product = {}
        for sale in self.sales:
            product = sale.get_product().get_product_id()
            quantity_sold = sale.get_quantity_sold()
            if product in sales_per_product:
                sales_per_product[product] += quantity_sold
            else:
                sales_per_product[product] = quantity_sold

        total_revenue = sum(sale.get_sale_price() * sale.get_quantity_sold() for sale in self.sales)

        delivery_times = [(sale.get_sale_date() - sale.get_client().get_registration_date()).days for sale in self.sales]
        average_delivery_time = sum(delivery_times) / len(delivery_times)

        most_demanded_product = max(sales_per_product, key=sales_per_product.get)

        return {
            "average_sales_per_client": average_sales_per_client,
            "total_sales_per_product": sales_per_product,
            "total_revenue": total_revenue,
            "average_delivery_time": average_delivery_time,
            "most_demanded_product": most_demanded_product }

  
    def generate_report(self):
        revenue = self._quantity_sold * self._sale_price
        report_date = self._sale_date 
        report = report(report_date, revenue)
        return report
    
class Invoice:
    def __init__(self, invoice_id, invoice_date=None, total_amount=None, products=None):
        self._invoice_id = invoice_id
        self._invoice_date = invoice_date
        self._total_amount = total_amount
        self._products = products
    
    def get_invoice_id(self):
        return self._invoice_id
    
    def set_invoice_id(self, invoice_id):
        self._invoice_id = invoice_id
    id=property(get_invoice_id,set_invoice_id)
    
    def get_invoice_date(self):
        return self._invoice_date
    
    def set_invoice_date(self, invoice_date):
        self._invoice_date = invoice_date
    date=property(get_invoice_date,set_invoice_date)
    
    def get_total_amount(self):
        return self._total_amount
    
    def set_total_amount(self, total_amount):
        self._total_amount = total_amount
    total_amount=property(get_total_amount,set_total_amount)
    
    def get_products(self):
        return self._products
    
    def set_products(self, products):
        self._products = products
    products=property(get_products,set_products)

class livraison:
    def __init__(self, shipping_nbr, client_name, cin, shipping_address, phone_nbr, order_nbr, order_date, 
                 central_dep_hr, central_arr_hr, worksite_arr_hr, worksite_dep_hr, product_id, quantity, 
                 product_type, vehicle, total_ttc, pompe1, pompe2, chauffeur, quantity_adju, livraison_date):
        self._shipping_nbr = shipping_nbr
        self._client_name = client_name
        self._cin = cin
        self._shipping_address = shipping_address
        self._phone_nbr = phone_nbr
        self._order_nbr = order_nbr
        self._order_date = order_date
        self._central_dep_hr = central_dep_hr
        self._central_arr_hr = central_arr_hr
        self._worksite_arr_hr = worksite_arr_hr
        self._worksite_dep_hr = worksite_dep_hr
        self._product_id = product_id
        self._quantity = quantity
        self._product_type = product_type 
        self._vehicle = vehicle
        self._total_ttc = total_ttc
        self._pompe1 = pompe1
        self._pompe2 = pompe2
        self._chauffeur = chauffeur
        self._quantity_adju = quantity_adju
        self._livraison_date = livraison_date


    def pompe1(self):
        return self._pompe1
    
    def Pompe1(self, pompe1):
        self._pompe1 = pompe1
        
    def pompe2(self):
        return self._pompe2
    
    def Pompe2(self, pompe2):
        self._pompe2 = pompe2

    def chauffeur(self):
        return self._chauffeur
    
    def Chauffeur(self, chauffeur):
        self._chauffeur = chauffeur
        
    def quantity_adju(self):
        return self._quantity_adju
    
    def Quantity_adju(self, quantity_adju):
        self._quantity_adju = quantity_adju

    def livraison_date(self):
        return self._livraison_date
    
    def Livraison_date(self, livraison_date):
        self._livraison_date = livraison_date
        
    def shipping_nbr(self):
        return self._shipping_nbr
    
    def Shipping_nbr(self, shipping_nbr):
        self._shipping_nbr = shipping_nbr    
         
    
    def client_name(self):
        return self._client_name
    
    def Client_name(self, client_name):
        self._client_name = client_name
        
    def cin(self):
        return self._cin
    
    def Cin(self, cin):
        self._cin = cin
        
    def shipping_address(self):
        return self._shipping_address
    
    def Shipping_address(self, shipping_address):
        self._shipping_address = shipping_address
        
    def phone_nbr(self):
        return self._phone_nbr
    
    def Phone_nbr(self, phone_nbr):
        self._phone_nbr = phone_nbr
        
    def order_nbr(self):
        return self._order_nbr
    
    def Order_nbr(self, order_nbr):
        self._order_nbr = order_nbr
        
    def order_date(self):
        return self._order_date
    
    def Order_date(self, order_date):
        self._order_date = order_date
        
    def central_dep_hr(self):
        return self._central_dep_hr
    
    def Central_dep_hr(self, central_dep_hr):
        self._central_dep_hr = central_dep_hr
        
    def central_arr_hr(self):
        return self._central_arr_hr
    
    def Central_arr_hr(self, central_arr_hr):
        self._central_arr_hr = central_arr_hr
        
    def worksite_arr_hr(self):
        return self._worksite_arr_hr
    
    def Worksite_arr_hr(self, worksite_arr_hr):
        self._worksite_arr_hr = worksite_arr_hr
        
    def worksite_dep_hr(self):
        return self._worksite_dep_hr
    
    def Worksite_dep_hr(self, worksite_dep_hr):
        self._worksite_dep_hr = worksite_dep_hr

    def product_id(self):
        return self._product_id

    def Product_id(self, product_id):
        self._product_id = product_id
    def quantity(self):
        return self._quantity
    
    def Quantity(self, quantity):
        self._quantity = quantity
        
    def product_type(self):
        return self._product_type
    
    def Product_type(self, product_type):
        self._product_type = product_type

    def vehicle(self):
        return self._vehicle
    
    def Vehicle(self, vehicle):
        self._vehicle = vehicle
        
    def total_ttc(self):
        return self._total_ttc
    
    def Total_ttc(self, total_ttc):
        self._total_ttc = total_ttc

    def afficher_details(self):
        print(f"Shipping Number : {self._shipping_nbr}")
        print(f"Client Name : {self._client_name}")
        print(f"CIN : {self._cin}")
        print(f"Shipping Address : {self._shipping_address}")
        print(f"Phone Number : {self._phone_nbr}")
        print(f"Order Number : {self._order_nbr}")
        print(f"Order Date : {self._order_date}")
        print(f"Central Departure Hour : {self._central_dep_hr}")
        print(f"Central Arrival Hour : {self._central_arr_hr}")
        print(f"Worksite Arrival Hour : {self._worksite_arr_hr}")
        print(f"Worksite Departure Hour : {self._worksite_dep_hr}")
        print(f"Product ID : {self._product_id}")
        print(f"Quantity : {self._quantity}")
        print(f"Product Type : {self._product_type}")
        print(f"Vehicle : {self._vehicle}")
        print(f"Total TTC : {self._total_ttc}")

    
#general functions to get objects by id    
def getClientById(id):
    for client in client_instances:
        if client.clt_id==id:
            return client
    return None

def getOrderById(id, index=False):
    for i,order in enumerate(order_instances):
        if order.order_id==id:
            if index:
                return i
            else:
                return order
    return None

def getProductById(id):
    for prod in product_instances:
        if prod.id==id:
            return prod
    return None

def getProductByDescription(description):
    for prod in product_instances:
        if prod.description==description:
            return prod
    return None



################## LOAD DATA ##################

# For Personne class
df1 = pd.read_csv("./class_personne.csv")
personne_instances = []
for index, row in df1.iterrows():
    personne_id = row['personne_id']
    name = row['name']
    address = row['address']
    phone_number = row['phone_number']
    personne = Personne(personne_id, name, address, phone_number)
    personne_instances.append(personne)
#for personne in personne_instances:
    #print(personne.id, personne.name, personne.address, personne.phone)

#For Supplier class
df2=pd.read_csv("./class_supplier.csv")
supplier_instances = []
for index, row in df2.iterrows():
    supplier_id = row['supplier_id']
    name = row['name']
    address = row['address']
    phone_number = row['phone_number']
    supplier = Supplier(supplier_id, name, address, phone_number)
    supplier_instances.append(supplier)
#for  supp in supplier_instances:
    #print(supp.id,supp.name,supp.address,supp.phone)

#For Employee class 
df3=pd.read_csv("./class_employee.csv")
employee_instances = []
for index, row in df3.iterrows():
    employee_id = row['employee_id']
    name = row['name']
    address = row['address']
    phone_number = row['phone_number']
    position=row['position']
    salary=row['salary']
    employee = Employee(employee_id, name, address, phone_number,position,salary)
    employee_instances.append(employee)
#for employee in employee_instances:
    #print(employee.id,employee.name,employee.address,employee.phone,employee.position,employee.salary)

#For Client class
df4=pd.read_csv("./class_client.csv")
client_instances = []
for index, row in df4.iterrows():
    client_id = row['ID']
    name = row['Name']
    address = row['Address']
    email_client = row['Email']
    phone_number = row['Phone Number']
    client = Client(client_id, name, address, phone_number, email_client)
    client_instances.append(client)
#for client in client_instances:
    #print(client.id,client.name,client.address,client.phone)

#For product class
df5=pd.read_csv("./class_product.csv")
product_instances = []

for index, row in df5.iterrows():
    product_id = row['product_id']
    description = row['description']
    price = row['price']
    quantity_in_stock = row['quantity_in_stock']
    product = Product(product_id, description, price, quantity_in_stock)
    product_instances.append(product)
tmp = [prod.description for prod in product_instances]
assert len(tmp)==len(set(tmp))
#2 produits ne peuvent avoir la même description

#for product in product_instances:
    #print(product.id, product.description, product.prix, product.stock)

#For Order class
df6=pd.read_csv("./class_order.csv")
order_instances = []

#csv columns : order_id, order_date, client, products, payment_type, price, price_paid, pompe, statut
for index, row in df6.iterrows():
    order_id = row['order_id']
    order_date = row['order_date']
    client = getClientById(int(row['client_id']))
    products = [(getProductById(int(id)), qty) for id, qty in (a.split(' : ') for a in row['products'].split(' | '))]
    payment_type=row['payment_type']
    price=float(row['price'])
    price_paid=float(row['price_paid'])
    pompe=True if row['pompe']=="True" else False
    statut=row['statut']
    order = Order(order_id, order_date, client, products, payment_type, price, price_paid, pompe, statut)
    order_instances.append(order)


#For Invoice class
df7=pd.read_csv("./class_invoice.csv")
invoice_instances = []

for index, row in df7.iterrows():
    invoice_id = row['invoice_id']
    invoice_date = row['invoice_date']
    total_amount = row['total_amount']
    products = row['product_ids']
    invoice = Invoice(invoice_id, invoice_date, total_amount, products)
    invoice_instances.append(invoice)

#for invoice in invoice_instances:
    #print(invoice.id,invoice.date , invoice.total_amount, invoice.products)

#For Warehouse class
df8=pd.read_csv("./class_warehouse.csv")
warehouse=Warehouse(10000,0) #on a choisit pour le moment capacity=10000
prd=0
for index, row in df8.iterrows():
    quant = row['quantity_in_stock']
    prd+=quant

#For Sale class
df9 = pd.read_csv("./class_sales.csv")
#print(df9)
sale_instances = []

for index, row in df9.iterrows():
    #sale_id = row['order_id']
    sale_date = row['Sale Date']
    product_id = row['Product ID']
    quantity_sold = row['Quantity Sold']
    sale_price = row['Sale Price']
    client_id = row['Client ID']
    sale = Sale(sale_date, product_id, quantity_sold, sale_price,client_id)
    sale_instances.append(sale)



#### USER INTERFACE 

##########################

# Sample data (you can load data from CSV as well)
clients_data = clients_data =pd.read_csv('./class_client.csv')
clients_data = clients_data.to_dict(orient='records')

# clients_data = [
#     {"ID": 1, "Name": "Alice", "Email": "alice@example.com", "Phone Number": "1234567890"},
#     {"ID": 2, "Name": "Bob", "Email": "bob@example.com", "Phone Number": "9876543210"},
#     {"ID": 3, "Name": "Charlie", "Email": "charlie@example.com", "Phone Number": "5678901234"},
#     {"ID": 4, "Name": "David", "Email": "david@example.com", "Phone Number": "8765432109"},
#     {"ID": 5, "Name": "Eve", "Email": "eve@example.com", "Phone Number": "4321098765"}
# ]

clients_data = clients_data =pd.read_csv('./class_client.csv')
clients_data = clients_data.to_dict(orient='records')

deleted_clients_data = pd.read_csv('./deleted_clients.csv')
deleted_clients_data= deleted_clients_data.to_dict(orient='records')

products_data = pd.read_csv('./class_product.csv')
products_data = products_data.to_dict(orient='records')

sales_data = pd.read_csv('./class_sales.csv')
sales_data = sales_data.to_dict(orient='records')


sales_data_displayed = sales_data

# Modification et lecture des fichiers csv

def csv_list(chemin):
    df_livraison=pd.read_csv(chemin, sep = ';', encoding = 'latin-1')
    list_livraison = df_livraison.values.tolist()
    for i in range(len(list_livraison)):
        list_livraison[i] = list_livraison[i][0].split(';')
    return list_livraison

def add_livraison_csv(new_line, chemin):
    with open(chemin, 'a', newline='') as file:
        csv_writer = csv.writer(file)
        csv_writer.writerow(new_line)
        
   

livraison_data = csv_list("./livraison.csv")




# Fonction pour se connecter
def login(event=None): #event=None to be able to login by pressing the Return key or using the button
    employee_name = username_entry.get()
    employee_password = password_entry.get()
    if employee_name == ADMIN_USERNAME:
        if employee_password == ADMIN_PASSWORD:
            login_window.destroy()
            create_main_window()
        else:
            # Show an error message for mot de passe incorrect
            messagebox.showerror("Erreur", "Mot de passe incorrect")
    else:
        # Show an error message for nom d'utilisateur incorrect
        messagebox.showerror("Erreur", "Nom d'utilisateur incorrect")        


def validate_id(action, index, value_if_allowed, prior_value, text, validation_type, trigger_type, widget_name):
    """Validate the ID field to allow only numeric input."""
    if text.isdigit() or text == "":
        return True
    else:
        return False
    
def validate_float(action, index, value_if_allowed, prior_value, text, validation_type, trigger_type, widget_name):
    """Validate the entry field to allow numeric inputs and . to allow conversion to float"""
    if text.isdigit() or text=="" or text ==".":
        return True
    elif text.replace(".","").isnumeric():
        return True
    else:
        return False


def get_next_order_id():
    existing_ids = [order.order_id for order in order_instances]
    new_id = 1
    while new_id in existing_ids:
        new_id += 1
    return new_id

def is_numeric_input(input_str):
    """Check if the input string is numeric."""
    return re.match(r'^\d+$', input_str) is not None

def is_price_input(input_str):
    """Check if the input string is numeric."""
    return re.match(r'^\d+\.\d{2}$', input_str) is not None

def is_float(string):
    if string.replace('.', '').isnumeric() and string.count('.')<2:
        return True
    else:
        return False

def validate_phone_number(action, index, value_if_allowed, prior_value, text, validation_type, trigger_type, widget_name):
    """Validate the phone number field to allow only numeric input."""
    if text.isdigit() or text == "":
        return True
    else:
        return False
    
def get_next_client_id():
    existing_ids = [client.clt_id for client in client_instances]
    new_id = 1
    while new_id in existing_ids:
        new_id += 1
    return new_id    

def refresh_client_ids():
    for index, client in enumerate(clients_data):
        client["ID"] = index + 1

def refresh_product_ids():
    for index, product in enumerate(products_data):
        product["ID"] = index + 1
        
def is_valid_phone_number(phone_number):
    # Validate the phone number using a regular expression
    phone_pattern = re.compile(r'^216[1-9]\d{0,10}$') #not starting by zero and contains 11 figures
    return phone_pattern.match(phone_number)

def is_valid_email(email):
    # Validate the email address using a regular expression
    email_pattern = re.compile(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$')
    return email_pattern.match(email)


def generate_new_supplier_id():
    # Retrieve existing supplier IDs from the CSV file
    csv_file_path = "./class_supplier.csv" # Replace with the actual path to your CSV file
    existing_ids = set()

def convert_docx_to_pdf(docx_filename, pdf_filename):
    # Create a Word application object
        
    current_directory = os.getcwd()

    # Create absolute paths from the relative paths
    absolute_input_path = os.path.join(current_directory, docx_filename)
    absolute_output_path = os.path.join(current_directory, pdf_filename)
    # Create a Word application object
    word = comtypes.client.CreateObject("Word.Application")
    
    # Open the Word document
    doc = word.Documents.Open(absolute_input_path)
    
    # Save the document as PDF
    doc.SaveAs(absolute_output_path, FileFormat=17)  # 17 represents the PDF format
    
    # Close the Word document and application
    doc.Close()
    word.Quit()   

def apply_deletion_on_startup():
    global deleted_clients_data
    try:
        with open('deletion_delay_weeks.txt', 'r') as f:
            delay_weeks = int(f.read())
    except (FileNotFoundError, ValueError):
        print("Fichier de délai de suppression non trouvé ou contenu invalide.")
        return

    # Calculer la date de suppression basée sur le délai en semaines
    deletion_date = datetime.date.today() - timedelta(weeks=delay_weeks)
    deletion_date = datetime.datetime.combine(deletion_date, datetime.datetime.min.time())

    # Filtrer deleted_clients_data
    updated_deleted_clients_data = [client for client in deleted_clients_data if datetime.datetime.strptime(client['DateSuppr'], '%d/%m/%Y') >= deletion_date]

    # Mettre à jour deleted_clients_data
    deleted_clients_data = updated_deleted_clients_data

    # Mettre à jour le fichier CSV
    df = pd.DataFrame(updated_deleted_clients_data)
    df.to_csv('deleted_clients.csv', index=False)

    print(f"Les enregistrements des clients supprimés ont été mis à jour, en conservant ceux après {deletion_date.strftime('%d/%m/%Y')}.")
    # try:
    #     # Étape 1: Lire la date de suppression du fichier
    #     with open('deletion_date.txt', 'r') as f:
    #         deletion_date_str = f.read()
    #     deletion_date = datetime.strptime(deletion_date_str, '%d/%m/%Y')#'%Y-%m-%d')
    # except (FileNotFoundError, ValueError):
    #     print("Date de suppression non définie ou fichier manquant. Opération annulée.")
    #     return

    # # Étape 2: Filtrer deleted_clients_data pour garder les clients dont la date de suppression est inférieure à la date de suppression
    # updated_deleted_clients_data = [client for client in deleted_clients_data if datetime.strptime(client['DateSuppr'], '%d/%m/%Y') < deletion_date]

    # # Étape 2: Charger les données des clients supprimés depuis le CSV
    # # try:
    # #     df = pd.read_csv('deleted_clients.csv', parse_dates=['DeletionDate'])
    # # except FileNotFoundError:
    # #     print("Fichier des clients supprimés introuvable.")
    # #     return

    # # Étape 3: Mettre à jour le fichier CSV
    # # Convertir updated_deleted_clients_data en DataFrame puis enregistrer dans le fichier CSV
    # df = pd.DataFrame(updated_deleted_clients_data)
    # df.to_csv('deleted_clients.csv', index=False)
    # # Étape 3 & 4: Filtrer pour garder les clients dont la date de suppression est inférieure à la date de suppression
    # # df_filtered = df[df['DeletionDate'] < deletion_date]

    # # Étape 4: Mettre à jour la variable deleted_clients_data avec les données filtrées
    # deleted_clients_data = updated_deleted_clients_data
    # # Étape 5: Écrire les données filtrées de retour dans le fichier CSV
    # # df_filtered.to_csv('deleted_clients.csv', index=False)
    # print("Les enregistrements des clients supprimés ont été mis à jour.")


# Création de la fenêtre principale
def create_main_window():
    apply_deletion_on_startup()
    window = tk.Tk()
    window.title("Gestion d'entreprise")
    window.geometry("1000x600+0+0")
    navbar = tk.Frame(window)
    navbar.pack()
    frame = tk.Frame(window)
    frame.pack()

    
        
    def Products():     
        Clear_widgets(frame)
        def display_products():
            if not products_tree.get_children():
                for product in products_data:
                    products_tree.insert("", tk.END, values=(product["product_id"], product["description"], product["price"], product["quantity_in_stock"], product["historique"]))

        def refresh_products_table():
            products_tree.delete(*products_tree.get_children())
            for product in products_data:
                products_tree.insert("", tk.END, values=(product["product_id"], product["description"], product["price"], product["quantity_in_stock"], product["historique"]))
        
        def double_click_product(event):
            selected_item = products_tree.focus()
            if selected_item:
                values = products_tree.item(selected_item, "values")
                if values:
                    product_id_entry.delete(0, tk.END)
                    description_entry.delete(0, tk.END)
                    price_entry.delete(0, tk.END)
                    quantity_entry.delete(0, tk.END)
                    historique_entry.delete(0, tk.END)
                    product_id_entry.insert(tk.END, values[0])
                    description_entry.insert(tk.END, values[1])
                    price_entry.insert(tk.END, values[2])
                    quantity_entry.insert(tk.END, values[3])
                    historique_entry.insert(tk.END, values[4])
                    
        def add_product():
            produit_id = product_id_entry.get()
            description = description_entry.get()
            price = price_entry.get()
            quantity_in_stock = quantity_entry.get()
            historique = historique_entry.get()
            if produit_id and description and price and quantity_in_stock and historique:
                if is_numeric_input(produit_id):
                    new_id = get_next_order_id()
                    
                    df = pd.read_csv("./class_product.csv")
                    size = df.shape[0] + 1

                    new_id = size
                    
                    df.loc[size] = [new_id, description, price, quantity_in_stock, historique]
                    df.to_csv('.class_product.csv', index=False)
                
                    new_product = Product(new_id, description, price, quantity_in_stock, historique)
                    product_instances.append(new_product)
                    products_tree.insert("", tk.END, values=(new_id, description, price, quantity_in_stock, historique))
                    product_id_entry.delete(0, tk.END)
                    description_entry.delete(0, tk.END)
                    price_entry.delete(0, tk.END)
                    quantity_in_stock.delete(0, tk.END)  # Set default value for Type de Transaction
                else:
                    messagebox.showerror("Erreur", "L'identifiant de client doit être une valeur numérique.")
            else:
                messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
                
        def delete_product():
            selected_item = products_tree.selection()
            if selected_item:
                item_id = products_tree.item(selected_item)["values"][0]
                df = pd.read_csv('./class_product.csv', sep = ',')
                df = df[df['product_id'] != item_id]
                df.to_csv('./class_product.csv', index=False)

                for product in products_data:
                    if product["product_id"] == item_id:
                        
                        
                        products_data.remove(product)
                        break
                products_tree.delete(selected_item)
                refresh_product_ids()
                refresh_products_table()
            else:
                messagebox.showwarning("Avertissement", "Veuillez sÃ©lectionner un produit Ã  supprimer.")
                
        def modify_product():
            selected_item = products_tree.selection()
            if selected_item:
                product_id = product_id_entry.get()
                description = description_entry.get()
                price = price_entry.get()
                quantity = quantity_entry.get()
                historique = historique_entry.get()

                if product_id and description and price and quantity and historique :
                    if is_numeric_input(product_id) and is_price_input(price) and is_numeric_input(quantity):
                        product_id = int(product_id)
                        selected_produit_id = products_tree.item(selected_item)["values"][0]
                        if selected_produit_id == product_id:
                            products_tree.item(selected_item, values=(product_id, description, price, quantity, historique))
                            
                            df = pd.read_csv("./class_product.csv")
                            colonne_index = 'product_id'
                            df = df.set_index(colonne_index)
                            
                            nouvelles_valeurs = {'description': description, 'price' : price, 'quantity_in_stock' : quantity, 'historique' : historique}
                            df.loc[product_id] = nouvelles_valeurs
                            df.reset_index(inplace = True)
                            
                            df.to_csv("./class_product.csv", index = False)
                            
                            messagebox.showinfo("Succès", "Commande modifiée avec succès.")
                        else:
                            messagebox.showerror("Erreur", "L'ID de la commande ne peut pas être modifié.")
                    else:
                        messagebox.showerror("Erreur", "L'identifiant de commande et l'identifiant de client doivent être des valeurs numériques.")
                else:
                    messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
            else:
                messagebox.showwarning("Avertissement", "Veuillez sélectionner une commande à modifier.")
    # Create the table to display products data
        titre_label = tk.Label(frame, text="Products", font=("Arial", 16))
        titre_label.pack(pady=5)
        
        columns_products = ("product_id", "description", "price", "quantity_in_stock", "historique")
        
        tree_frame = tk.Frame(frame)
        tree_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical")
        scrollbar.pack(side='right', fill='y')
        
        style = ttk.Style()
        style.configure('Treeview', rowheight=25)
        products_tree = ttk.Treeview(tree_frame, columns=columns_products, show="headings", style='Custom.Treeview')
        
        scrollbar.config(command=products_tree.yview)  
        
        for col in columns_products:
            products_tree.heading(col, text=col)
            products_tree.column(col, width=5)

        products_tree.pack(fill=tk.BOTH, expand=True, pady=5)
        products_tree.bind("<Double-1>", double_click_product)
        
        # Add the labels and input fields for adding/modifying an order
        input_frame_products = tk.Frame(frame)
        input_frame_products.pack()

        product_id_label = tk.Label(input_frame_products, text="ID du produit :")
        product_id_label.pack(side=tk.LEFT, padx=5)
        product_id_entry = tk.Entry(input_frame_products)
        product_id_entry.pack(side=tk.LEFT, padx=5)

        description_label = tk.Label(input_frame_products, text="Description :")
        description_label.pack(side=tk.LEFT, padx=5)
        description_entry = tk.Entry(input_frame_products)
        description_entry.pack(side=tk.LEFT, padx=5)


        price_label = tk.Label(input_frame_products, text="Prix :")
        price_label.pack(side=tk.LEFT, padx=5)
        price_entry = tk.Entry(input_frame_products)
        price_entry.pack(side=tk.LEFT, padx=5)

        quantity_label = tk.Label(input_frame_products, text="Quantité :")
        quantity_label.pack(side=tk.LEFT, padx=5)
        quantity_entry = tk.Entry(input_frame_products)
        quantity_entry.pack(side=tk.LEFT, padx=5)

        historique_label = tk.Label(input_frame_products, text="Historique :")
        historique_label.pack(side=tk.LEFT, padx=5)
        historique_entry = tk.Entry(input_frame_products)
        historique_entry.pack(side=tk.LEFT, padx=5)
        
        button_frame_products = tk.Frame(frame)
        button_frame_products.pack(pady=5)
        
        display_products = ctk.CTkButton(button_frame_products, text="Liste des produits", command=display_products)
        display_products.pack(side=tk.LEFT, padx=5)
        # Add the buttons for adding/modifying a product


        add_button_products = tk.Button(button_frame_products, text="Ajouter Produit", command=add_product)
        add_button_products.pack(side=tk.LEFT, padx=5)

        delete_button_orders = tk.Button(button_frame_products, text="Supprimer Produit", command=delete_product)
        delete_button_orders.pack(side=tk.LEFT, padx=5)

        modify_button_orders = tk.Button(button_frame_products, text="Modifier Produit", command=modify_product)
        modify_button_orders.pack(side=tk.LEFT, padx=5)

        
    def Supplier():
        Clear_widgets(frame)
        def refresh_suppliers():
            # Clear old data from the display
            suppliers_tree.delete(*suppliers_tree.get_children())

            # Retrieve supplier data from the CSV file
            csv_file_path = "./class_supplier.csv" 

            with open(csv_file_path, newline='', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile)
                suppliers_data = list(reader)

            for supplier in suppliers_data:
                suppliers_tree.insert("", tk.END, values=(
                    supplier["supplier_id"], supplier["name"], supplier["address"],
                    supplier["phone_number"], supplier["contact_person"], supplier.get("email", "")
                ))

            # Clear the entry fields
            for entry in entries:
                entry.delete(0, tk.END)
        
        def double_click_supplier(event):
            selected_item = suppliers_tree.focus()
            if selected_item:
                values = suppliers_tree.item(selected_item, "values")
                if values:
                    # Clear the existing data from the specific entry variables
                    supplier_id_entry_supplier.delete(0, tk.END)
                    name_entry_supplier.delete(0, tk.END)
                    address_entry_supplier.delete(0, tk.END)
                    phone_entry_supplier.delete(0, tk.END)
                    contact_person_entry_supplier.delete(0, tk.END)
                    email_entry_supplier.delete(0, tk.END)

                    # Insert the values into the specific entry variables
                    supplier_id_entry_supplier.insert(tk.END, values[0])
                    name_entry_supplier.insert(tk.END, values[1])
                    address_entry_supplier.insert(tk.END, values[2])
                    phone_entry_supplier.insert(tk.END, values[3])
                    contact_person_entry_supplier.insert(tk.END, values[4])
                    email_entry_supplier.insert(tk.END, values[5])
        
        def add_supplier():
            try:
                global name_entry_supplier, address_entry_supplier, phone_entry_supplier, contact_person_entry_supplier, email_entry_supplier

                # Retrieve values from the entry fields using specific entry variables
                new_name = name_entry_supplier.get()
                new_address = address_entry_supplier.get()
                new_main_number_value = phone_entry_supplier.get()
                new_contact_person = contact_person_entry_supplier.get()
                new_email = email_entry_supplier.get()

                if len(new_main_number_value) == 11: 
                    new_main_number = new_main_number_value  # if phone lenght = 11, no need to add 216
                else:
                    new_main_number = '216' + new_main_number_value  # add 216 in other cases


                # Validate the phone number
                if not is_valid_phone_number(new_main_number):
                    messagebox.showerror("Erreur", "Le numéro de téléphone doit commencer par 216 et contenir de 1 à 9 chiffres.")
                    return  # Exit the function if the phone number is not valid

                # Validate the email address
                if not is_valid_email(new_email):
                    messagebox.showerror("Erreur", "L'adresse e-mail n'est pas valide.")
                    return  # Exit the function if the email address is not valid

                # Generate a new ID automatically
                new_id = generate_new_supplier_id()

                # Create a new supplier object
                new_supplier = Supplier(new_id, new_name, new_address, new_main_number, new_contact_person, new_email)

                # Insert the new supplier data into the CSV file
                csv_file_path = "./class_supplier.csv"

                with open(csv_file_path, mode='a', newline='', encoding='utf-8') as csv_file:
                    csv_writer = csv.writer(csv_file, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
                    csv_writer.writerow([
                        new_supplier._personne_id,  # Assuming this is the attribute for ID
                        new_supplier._name,
                        new_supplier._address,
                        new_supplier._phone_number,  # Assuming this is the attribute for phone number
                        new_supplier.contact_person,
                        new_supplier.email
                    ])

                # Update the display
                refresh_suppliers()
                messagebox.showinfo("Succès", "Fournisseur ajouté avec succès.")

            except Exception as e:
                messagebox.showerror("Erreur", f"Une erreur s'est produite : {str(e)}")
                
                
        def delete_supplier():
            # Retrieve the supplier ID to delete
            supplier_id_to_delete = int(supplier_id_entry_supplier.get())  # Assume the ID to delete is entered in the supplier ID field

            # Read the existing data from the CSV file
            csv_file_path = "./class_supplier.csv"

            with open(csv_file_path, mode='r', newline='', encoding='utf-8') as csv_file:
                csv_reader = csv.reader(csv_file)
                rows = list(csv_reader)

            deleted_supplier = None

            # Find and delete the supplier in the CSV data
            for idx, row in enumerate(rows):
                if idx > 0:  # Skip the header row
                    current_supplier_id = int(row[0])
                    if current_supplier_id == supplier_id_to_delete:
                        deleted_supplier = rows.pop(idx)
                        break

            if deleted_supplier:
                # Update the CSV file with the modified data
                with open(csv_file_path, mode='w', newline='', encoding='utf-8') as csv_file:
                    csv_writer = csv.writer(csv_file, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
                    csv_writer.writerows([rows[0]] + rows[1:])  # Write the header row and updated data

                # Update the display
                refresh_suppliers()

                # Clear the entry fields for the deleted supplier from the interface using specific entry variables
                supplier_id_entry_supplier.delete(0, tk.END)
                name_entry_supplier.delete(0, tk.END)
                address_entry_supplier.delete(0, tk.END)
                phone_entry_supplier.delete(0, tk.END)
                contact_person_entry_supplier.delete(0, tk.END)
                email_entry_supplier.delete(0, tk.END)

                # Inform about the deletion
                messagebox.showinfo("Suppression", f"Fournisseur {deleted_supplier[1]} supprimé avec succès.")
            else:
                messagebox.showerror("Erreur", f"Aucun fournisseur trouvé avec l'ID {supplier_id_to_delete}.")
                
                
        def modify_supplier():
            # Get the ID of the supplier to modify
            id_to_modify = int(supplier_id_entry_supplier.get())

            # Read the existing data from the CSV file
            csv_file_path = "./class_supplier.csv"

            with open(csv_file_path, mode='r', newline='', encoding='utf-8') as csv_file:
                csv_reader = csv.reader(csv_file)
                rows = list(csv_reader)

            existing_supplier = None
            idx_to_modify = None

            # Find the supplier in the CSV data
            for idx, row in enumerate(rows):
                if idx > 0:  # Skip the header row
                    current_supplier_id = int(row[0])
                    if current_supplier_id == id_to_modify:
                        existing_supplier = row
                        idx_to_modify = idx
                        break

            if existing_supplier is not None:
                # Get the updated values from the interface
                new_name = name_entry_supplier.get()
                new_address = address_entry_supplier.get()
                new_contact_person = contact_person_entry_supplier.get()
                new_email = email_entry_supplier.get()
                new_phone_number = phone_entry_supplier.get()

                # Validate the updated phone number if it's not empty
                if new_phone_number and not is_valid_phone_number(new_phone_number):
                    messagebox.showerror("Erreur", "Le numéro de téléphone doit commencer par 216 et contenir de 1 à 9 chiffres.")
                    return  # Exit the function if the phone number is not valid

                # Validate the updated email if it's not empty
                if new_email and not is_valid_email(new_email):
                    messagebox.showerror("Erreur", "L'adresse e-mail n'est pas valide.")
                    return  # Exit the function if the email address is not valid

                # Update the supplier details in the CSV data
                existing_supplier[1] = new_name
                existing_supplier[2] = new_address
                existing_supplier[4] = new_contact_person

                # Update phone number only if it's not empty and valid
                if new_phone_number:
                    existing_supplier[3] = new_phone_number

                # Update email only if it's not empty and valid
                if new_email:
                    existing_supplier[5] = new_email

                # Update the CSV file with the modified data
                with open(csv_file_path, mode='w', newline='', encoding='utf-8') as csv_file:
                    csv_writer = csv.writer(csv_file, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
                    csv_writer.writerows([rows[0]] + rows[1:])  # Write the header row and updated data

                # Refresh the display
                refresh_suppliers()
                messagebox.showinfo("Succès", "Détails du fournisseur mis à jour avec succès.")
            else:
                messagebox.showerror("Erreur", "Aucun fournisseur trouvé avec l'ID spécifié.")


        # Create the table to display suppliers data
        titre_label = tk.Label(frame, text="Fournisseurs", font=("Arial", 16))
        titre_label.pack(pady=5)

        columns_suppliers = ("Supplier ID", "Name", "Address", "Phone Number", "Contact Person", "Email")

        tree_frame_suppliers = tk.Frame(frame)
        tree_frame_suppliers.pack(fill=tk.BOTH, expand=True, pady=10)

        scrollbar_suppliers = ttk.Scrollbar(tree_frame_suppliers, orient="vertical")
        scrollbar_suppliers.pack(side='right', fill='y')

        style_suppliers = ttk.Style()
        style_suppliers.configure('Treeview', rowheight=25)

        suppliers_tree = ttk.Treeview(tree_frame_suppliers, columns=columns_suppliers, show="headings", style='Custom.Treeview')

        scrollbar_suppliers.config(command=suppliers_tree.yview)

        for col in columns_suppliers:
            suppliers_tree.heading(col, text=col)
            suppliers_tree.column(col, width=150)

        suppliers_tree.pack(fill=tk.BOTH, expand=True, pady=10)
        suppliers_tree.bind("<Double-1>", double_click_supplier)

        # Add the labels and input fields for adding/modifying a supplier
        input_frame_suppliers = tk.Frame(frame)
        input_frame_suppliers.pack()

        fields = ["ID du Fournisseur", "Nom du Fournisseur", "Adresse du Fournisseur",
                "Numéro de téléphone du Fournisseur", "Nom du Contact", "Email du Fournisseur"]

        entries = []

        for idx, field_text in enumerate(fields):
            row_num = idx // 2
            col_num = idx % 2 * 2 + 1

            label = tk.Label(input_frame_suppliers, text=f"{field_text} :")
            label.grid(row=row_num, column=col_num - 1, padx=5, pady=5, sticky="e")

            entry = tk.Entry(input_frame_suppliers)
            entry.grid(row=row_num, column=col_num, padx=5, pady=5)
            entries.append(entry)

        # Access the entry fields using the list entries with specific names
        supplier_id_entry_supplier = entries[0]
        name_entry_supplier = entries[1]
        address_entry_supplier = entries[2]
        phone_entry_supplier = entries[3]
        contact_person_entry_supplier = entries[4]
        email_entry_supplier = entries[5]


        # Add the buttons for adding/modifying a supplier
        button_frame_suppliers = tk.Frame(frame)
        button_frame_suppliers.pack(pady=10)
        
        display_refresh_suppliers = tk.Button(button_frame_suppliers, text="Liste des fournisseurs", command=refresh_suppliers)
        display_refresh_suppliers.pack(side=tk.LEFT, padx=5)

        add_button_suppliers = tk.Button(button_frame_suppliers, text="Ajouter Fournisseur", command=add_supplier)
        add_button_suppliers.pack(side=tk.LEFT, padx=5)

        delete_button_suppliers = tk.Button(button_frame_suppliers, text="Supprimer Fournisseur", command=delete_supplier)
        delete_button_suppliers.pack(side=tk.LEFT, padx=5)

        modify_button_suppliers = tk.Button(button_frame_suppliers, text="Modifier Fournisseur", command=modify_supplier)
        modify_button_suppliers.pack(side=tk.LEFT, padx=5)



        button_refresh_suppliers = tk.Button(button_frame_suppliers, text="Rafraîchir", command=refresh_suppliers)
        button_refresh_suppliers.pack(side=tk.LEFT, padx=5)

    def Sales():
        Clear_widgets(frame)
        tk.Label(frame,text="Sales").pack()
        
        def add_sale():
            #sale_id = sale_id_entry.get()
            date = sale_date_entry.get()
            prod_id = product_id_entry.get()
            quant = quantity_sold_entry.get()
            price = sale_price_entry.get()
            client_id = client_id_entry.get()
            #print(type(price))
            
            if (#sale_id and 
            date and prod_id and quant and price and client_id):
                if (#is_numeric_input(sale_id) and 
                    is_numeric_input(prod_id) and is_numeric_input(quant) 
                and is_price_input(price) and is_numeric_input(client_id)):
                    
                    df = pd.read_csv('./class_sales.csv')
                    df2 = pd.read_csv('./deleted_sales.csv')
                    size = df.shape[0] + df2.shape[0] + 1

                    print(df.shape[0],df2.shape[0],size)
                    new_id = size

                    print(size)

                    df.loc[size] = [new_id, date, prod_id, quant, price, client_id]
                    df.to_csv('./class_sales.csv', index=False)

                    new_sale = {
                        "Sale ID": new_id, #/!\ nouvel attribut 
                        "Sale Date": date,
                        "Product ID": prod_id,
                        "Quantity Sold": quant,
                        "Sale Price": price,
                        "Client ID": client_id
                    }

                    sales_data.append(new_sale)
                    #sales_data_displayed.append(new_sale)

                    sales_tree.insert("", tk.END, values=(
                        new_sale["Sale ID"], new_sale["Sale Date"], 
                        new_sale["Product ID"], new_sale["Quantity Sold"],
                        new_sale["Sale Price"],new_sale["Client ID"]
                    ))
                    
                    sale_id_entry.delete(0, tk.END)
                    sale_date_entry.delete(0, tk.END)
                    product_id_entry.delete(0, tk.END)
                    quantity_sold_entry.delete(0, tk.END)
                    sale_price_entry.delete(0, tk.END)
                    client_id_entry.delete(0, tk.END)
                    #type_transaction_var.set("CHEQUES")  # Set default value for Type de Transaction
                    #statut_var.set("PAYE")  # Set default value for Statut
                else:
                    message=""
                    #if not is_numeric_input(sale_id) : message+="id de la vente;"
                    if not is_numeric_input(prod_id) : message+="id du produit;"
                    if not is_numeric_input(quant) : message+="quantité vendue;"
                    if not is_price_input(price) : 
                        message+="prix de vente;"
                        print(type(price))
                    if not is_numeric_input(client_id) : message+="id du client;"            
                    messagebox.showerror("Erreur", 
                                        f"L'id de la vente/l'id du produit/la quantité vendue/le prix de vente/l'identifiant de client\ndoivent être des valeurs numériques.\nErreur sur : {message}")
            else:
                messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")

        def delete_sale():
            selected_item = sales_tree.selection()
            if selected_item:
                sale_id = sales_tree.item(selected_item)["values"][0]
        
                df = pd.read_csv('./class_sales.csv', sep = ',')
                deleted_sale = df[df['Sale ID'] == sale_id]
                
                # add the deleted client to the deleted_clients csv file for archive
                with open('./deleted_sales.csv', 'a', newline='') as f:
                    deleted_sale.to_csv(f, header=f.tell()==0, index=False)
                    
                df = df[df['Sale ID'] != sale_id]
                df.to_csv('./class_sales.csv', index=False)
            
                #for sale in sales_data_displayed
                for sale in sales_data:
                    if sale["Sale ID"] == sale_id:
                        sales_data.remove(sale)
                        #sales_data_displayed.remove(sale)             # on ne devrait pas supprimer des données
                        break
                sales_tree.delete(selected_item)
                #refresh_sale_ids()                         #A VOIR SI A IMPLEMENTER
            else:
                messagebox.showwarning("Avertissement", "Veuillez sélectionner une vente à supprimer.")

        def display_sales():
            if not sales_tree.get_children():
                #for sale in sales_data_displayed:
                for sale in sales_data:
                    #print(sale)
                    sales_tree.insert("", tk.END, values=(
                        sale["Sale ID"],sale["Sale Date"], sale["Product ID"], sale["Quantity Sold"], sale["Sale Price"],sale["Client ID"]
                    ))

        def modify_sale():
            selected_item = sales_tree.selection()
            if selected_item:
                sale_id = sale_id_entry.get()
                sale_date = sale_date_entry.get()
                product_id = product_id_entry.get()
                quantity_sold = quantity_sold_entry.get()
                sale_price = sale_price_entry.get()
                client_id = client_id_entry.get()

                if sale_id and sale_date and product_id and quantity_sold and sale_price and client_id:
                    if is_numeric_input(sale_id) and is_numeric_input(product_id) and is_price_input(sale_price) and is_numeric_input(client_id):
                        sale_id=int(sale_id)
                        selected_sale_id = sales_tree.item(selected_item)["values"][0]
                        selected_product_id = sales_tree.item(selected_item)["values"][2]
                        selected_client_id = sales_tree.item(selected_item)["values"][5]

                        #if selected_sale_id == sale_id:
                        

                        print(type(selected_sale_id),type(sale_id), type(selected_client_id),type(client_id))
                        if selected_sale_id == int(sale_id) and selected_product_id == int(product_id) and selected_client_id == int(client_id):
                            df = pd.read_csv('./class_sales.csv')
                            colonne_index = 'Sale ID'
                            df = df.set_index(colonne_index)
                                
                            nouvelles_valeurs = {'Sale ID': sale_id, 'Sale Date': sale_date, 'Product ID' : product_id, 'Quantity Sold' : quantity_sold, 'Sale Price' : sale_price, 'Client ID' : client_id}
                            df.loc[sale_id] = nouvelles_valeurs #product_id
                            df.reset_index(inplace = True)
                                
                            df.to_csv('./class_sales.csv', index = False)

                            sales_tree.item(selected_item, values=(sale_id, sale_date, product_id, quantity_sold,sale_price,client_id))
                            messagebox.showinfo("Succès", "Vente modifiée avec succès.")
                        else:
                            messagebox.showerror("Erreur", "L'id de la vente et/ou l'id du client ne peut/peuvent pas être modifié.s.")
                            #mais effectue un ajout au final
                    else:
                        messagebox.showerror("Erreur", "l'id de vente/l'id de produit/la quantité vendue/l'id du client doivent être des valeurs numériques.")
                else:
                    messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
            else:
                messagebox.showwarning("Avertissement", "Veuillez sélectionner une vente à modifier.")            

        def double_click_sale(event):
            selected_item = sales_tree.focus()
            if selected_item:
                values = sales_tree.item(selected_item, "values")
                if values:
                    #print(values)
                    sale_id_entry.delete(0, tk.END)
                    sale_date_entry.delete(0, tk.END)
                    product_id_entry.delete(0, tk.END)
                    quantity_sold_entry.delete(0, tk.END)
                    sale_price_entry.delete(0, tk.END)
                    client_id_entry.delete(0, tk.END)
                    
                    sale_id_entry.insert(tk.END, values[0])
                    sale_date_entry.insert(tk.END, values[1])
                    product_id_entry.insert(tk.END, values[2])
                    quantity_sold_entry.insert(tk.END, values[3])
                    sale_price_entry.insert(tk.END, values[4])
                    client_id_entry.insert(tk.END, values[5])

    # Create the table to display sales data                
        titre_label = tk.Label(frame, text="Ventes", font=("Arial", 16))
        titre_label.pack(pady=5)
        
        columns_sales = ("Sale ID","Sale Date","Product ID","Quantity Sold","Sale Price","Client ID")
        
        tree_frame = tk.Frame(frame)
        tree_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical")
        scrollbar.pack(side='right', fill='y')

        style = ttk.Style()
        style.configure('Treeview', rowheight=25)
        
        sales_tree = ttk.Treeview(tree_frame, columns=columns_sales, show="headings", style='Custom.Treeview')

        scrollbar.config(command=sales_tree.yview)

        for col in columns_sales:
            sales_tree.heading(col, text=col)
            sales_tree.column(col, width=5)

        sales_tree.pack(fill=tk.BOTH, expand=True, pady=5)
        sales_tree.bind("<Double-1>", double_click_sale)
        
    # Add the labels and input fields for adding/modifying a sale
        input_frame_sales = tk.Frame(frame)
        input_frame_sales.pack()

        sale_id_label = tk.Label(input_frame_sales, text="ID de la vente :")
        sale_id_label.pack(side=tk.LEFT, padx=5)
        sale_id_entry = tk.Entry(input_frame_sales, validate="key")
        sale_id_entry.pack(side=tk.LEFT, padx=5)

        sale_date_label = tk.Label(input_frame_sales, text="Date de vente:")
        sale_date_label.pack(side=tk.LEFT, padx=5)
        sale_date_entry = tk.Entry(input_frame_sales, validate="key")
        """checker si la sale_date est valide"""
        #sale_date_entry.config(valsale_dateatecommand=(root.register(valisale_date_id), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
        sale_date_entry.pack(side=tk.LEFT, padx=5)

        product_id_label = tk.Label(input_frame_sales, text="ID du Produit :")
        product_id_label.pack(side=tk.LEFT, padx=5)
        product_id_entry = tk.Entry(input_frame_sales, validate="key")
        product_id_entry.config(validatecommand=(frame.register(validate_id), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
        product_id_entry.pack(side=tk.LEFT, padx=5)

        quantity_sold_label = tk.Label(input_frame_sales, text="Quantité vendue :")
        quantity_sold_label.pack(side=tk.LEFT, padx=5)
        quantity_sold_entry = tk.Entry(input_frame_sales, validate="key")
        "checker si la quantité est valide : 1. en tant qu'int positif;"               #check
        "2. en considérant le stock"
        quantity_sold_entry.config(validatecommand=(frame.register(validate_id), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
        quantity_sold_entry.pack(side=tk.LEFT, padx=5)

        sale_price_label = tk.Label(input_frame_sales, text="Prix de vente:")
        sale_price_label.pack(side=tk.LEFT, padx=5)
        sale_price_entry = tk.Entry(input_frame_sales, validate="key")
        """checker si le prix est valide: 1. en tant que float"""
        sale_price_entry.config(validatecommand=(frame.register(is_price_input), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
        sale_price_entry.pack(side=tk.LEFT, padx=5)

        client_id_label = tk.Label(input_frame_sales, text="ID du client:")
        client_id_label.pack(side=tk.LEFT, padx=5)
        client_id_entry = tk.Entry(input_frame_sales, validate="key")
        client_id_entry.config(validatecommand=(frame.register(validate_id), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
        """checker si le client existe bel et bien dans la database"""
        client_id_entry.pack(side=tk.LEFT, padx=5)

    # Add the buttons for adding/modifying a sale
        button_frame_sales = tk.Frame(frame)
        button_frame_sales.pack(pady=5)

        display_sales = ctk.CTkButton(button_frame_sales, text="Liste des ventes", command=display_sales)
        display_sales.pack(side=tk.LEFT, padx=5)

        add_button_sales = tk.Button(button_frame_sales, text="Ajouter Vente", command=add_sale)
        add_button_sales.pack(side=tk.LEFT, padx=5)

        delete_button_sales = tk.Button(button_frame_sales, text="Supprimer Vente", command=delete_sale)
        delete_button_sales.pack(side=tk.LEFT, padx=5)

        modify_button_sales = tk.Button(button_frame_sales, text="Modifier Vente", command=modify_sale)
        modify_button_sales.pack(side=tk.LEFT, padx=5)
        
    def Deliveries():
        Clear_widgets(frame)
        tk.Label(frame,text="Deliveries").pack()
        
        def supress_livraison():
    
            def on_no():
                root.destroy()

            def real_suppress():

                selected_item = livraison_tree.selection()
                shp_nbr = livraison_tree.item(selected_item)["values"][0]
                
                with open('livraison_data.csv', 'r', newline='') as file:
                    reader = csv.reader(file, delimiter=',')
                    lines = list(reader)

                # Find and remove the specified ID from the first column
                updated_lines = [','.join(row) for row in lines if str(shp_nbr) != str(row[0])]

                # Write the updated values back to the CSV file
                with open('livraison_data.csv', 'w', newline='') as file:
                    file.write('\n'.join(updated_lines))
                display_livraison()   
                root.destroy()

            root = tk.Tk()
            root.title("Confirmation")
            root.geometry('350x150')

            label = tk.Label(root, text="Voulez vous vraiment supprimer cette livraison ?")
            label.pack(pady=10)

            yes_button = tk.Button(root, text="OUI", command=real_suppress)
            yes_button.pack(side=tk.LEFT, padx=95)

            no_button = tk.Button(root, text="NON", command=on_no)
            no_button.pack(side=tk.LEFT, padx=5) 
        

        def double_click_livraison(event):
            selected_item = livraison_tree.focus()
            if selected_item:
                values = livraison_tree.item(selected_item, "values")
                if values:
                    n_entry.delete(0, tk.END)
                    name_entry.delete(0, tk.END)
                    adress_entry.delete(0, tk.END)
                    prod_entry.delete(0, tk.END)
                    n_entry.insert(tk.END, values[0])
                    name_entry.insert(tk.END, values[1])
                    adress_entry.insert(tk.END, values[2])
                    prod_entry.insert(tk.END, values[3])
        
        def modify_livraison():
            selected_item = livraison_tree.selection()
            shp_nbr = livraison_tree.item(selected_item)["values"][0]
            mod_line=[]
            with open('livraison_data.csv', 'r', newline='') as file:
                    reader = csv.reader(file, delimiter=',')
                    lines = list(reader)
            for row in lines :
                if str(row[0])==str(shp_nbr):
                    mod_line=row
            n_entry.insert(tk.END, row[0])
            page_livraison(2)

        def replace_placeholder_with_value(run, value):
            run.clear()  # Clear existing content
            run.add_text(value)  # Add the new text

        def remplir_template(livraison):
            # Charger le document Word existant (le modèle)
            doc = Document("TEMPLATE_Bon_de_livraison.docx")

            # Remplacer les espaces réservés dans les paragraphes
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if "<<INSERT_ORDER_NUMBER_HERE>>" in run.text:
                        replace_placeholder_with_value(run, str(livraison.shipping_nbr()))
                    if "<<INSERT_DATE_HERE>>" in run.text:
                        replace_placeholder_with_value(run, str(livraison.livraison_date()))
                    if "INSERTCLIENTCINMEF_HERE" in run.text:
                        replace_placeholder_with_value(run, str(livraison.cin()))
                    if "<<INSERT_CLIENT_NAME_HERE>>" in run.text:
                        replace_placeholder_with_value(run, livraison.client_name())
                    if "DD" in run.text:
                        replace_placeholder_with_value(run, str(livraison.order_date()))
                    if "TEL" in run.text:
                        replace_placeholder_with_value(run, str(livraison.phone_nbr()))
                    if "ORDERNUM" in run.text:
                        replace_placeholder_with_value(run, str(livraison.order_nbr()))
                    if "UU" in run.text:
                        replace_placeholder_with_value(run, str(livraison.shipping_address()))

            # Remplacer les espaces réservés dans les cellules des tableaux
            for table in  doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for run in cell.paragraphs[0].runs:
                            if "HH" in run.text:
                                replace_placeholder_with_value(run, str(livraison.central_dep_hr()))
                            if "RR" in run.text:
                                replace_placeholder_with_value(run, str(livraison.central_arr_hr()))
                            if "AA" in run.text:
                                replace_placeholder_with_value(run, str(livraison.worksite_arr_hr()))
                            if ".Chanti" in run.text:
                                replace_placeholder_with_value(run, str(livraison.worksite_dep_hr()))
                            if "INSERT_Type_HERE" in run.text:
                                replace_placeholder_with_value(run, livraison.product_type())
                            if "INSERT_Qte_HERE" in run.text:
                                replace_placeholder_with_value(run, str(livraison.quantity()))
                            if "TT" in run.text:
                                replace_placeholder_with_value(run, str(livraison.pompe1()))
                            if "INSERT_Tickl_HERE" in run.text:
                                replace_placeholder_with_value(run, str(livraison.pompe2()))
                            if "INSERT_Nature_HERE" in run.text:
                                replace_placeholder_with_value(run, livraison.product_type())
                            if "INSERT_Nature_HERE" in run.text:
                                replace_placeholder_with_value(run, livraison.product_type())
                            if "INSERT_Tel_HERE" in run.text:
                                replace_placeholder_with_value(run, str(livraison.quantity_adju()))
                            if "INSERT_Mat_HERE" in run.text:
                                replace_placeholder_with_value(run, livraison.vehicle())
                            if "INSERT_Chauff_HERE" in run.text:
                                replace_placeholder_with_value(run, livraison.chauffeur())
                            if "INSERT_Total_HERE" in run.text:
                                replace_placeholder_with_value(run, str(livraison.total_ttc()))

            out = str(livraison.shipping_nbr()) + ".docx"

            # Sauvegarder le document rempli
            doc.save(out)


        def page_livraison(R):

            def mod_livraison(shp_nbr):
            
                shipping_nbr = n_entry.get()
                nouvelles_donnees = []
                
                if(str(shp_nbr)==str(shipping_nbr)):
                    client_name = name_entry.get()
                    cin = cin_entry.get()
                    shipping_address = adress_entry.get()
                    phone_nbr = phone_nbr_entry.get()
                    order_nbr = order_entry.get()
                    order_date = order_date_entry.get_date() 
                    order_date_l= order_date_l_entry.get_date() 
                    central_dep_hr = central_dep_hr_combobox.get()
                    central_arr_hr = central_arr_hr_combobox.get()
                    worksite_dep_hr = worksite_dep_hr_combobox.get()
                    worksite_arr_hr = worksite_arr_hr_combobox.get()    
                    product_id = prod_entry.get()
                    adjuvant=adjuvant_combobox.get()
                    quantity = quantite_entry.get()
                    quantity_a = quantity_a_entry.get()
                    vehicle = matricule_entry.get()
                    total_ttc = total_ttc_entry.get()
                    chauffeur = chauffeur_combobox.get()
                    variable1 = variable1_combobox.get() 


                    if (shipping_nbr and client_name and cin and shipping_address and phone_nbr and order_nbr and order_date and central_dep_hr and central_arr_hr and
                    worksite_dep_hr and worksite_arr_hr and product_id and quantity and vehicle and total_ttc and chauffeur and adjuvant and order_date_l and quantity_a and variable1):
                        
                        with open('livraison_data.csv', 'r', newline='') as file:
                                reader = csv.reader(file, delimiter=',')
                                lines = list(reader)
                        for row in lines :
                            if str(row[0])==str(shp_nbr):
                                row[1]=client_name
                                row[2]=cin
                                row[3]=shipping_address
                                row[4]=phone_nbr
                                row[5]=order_nbr
                                row[6]=order_date
                                row[7]=central_dep_hr
                                row[8]=central_arr_hr
                                row[9]=worksite_dep_hr
                                row[10]=worksite_arr_hr
                                row[11]=product_id
                                row[12]=total_ttc
                                row[13]=adjuvant
                                row[14]=vehicle
                                row[15]=quantity
                                if variable1 == "OUI":
                                    row[16] ="X"
                                else:row[16]=" " 
                                if row[16]=='X': row[17]=' ' 
                                else : row[17]='X'
                                row[18]=chauffeur
                                row[19]=quantity_a
                                row[20]=order_date_l
                                nouvelles_donnees.append(row)
                            else:
                                nouvelles_donnees.append(row)
                        with open('livraison_data.csv', 'w', newline='') as fichier_csv:
                            ecrivain_csv = csv.writer(fichier_csv)
                            ecrivain_csv.writerows(nouvelles_donnees)
                        root.withdraw()   
                        root.after(100, root.destroy)
                        display_livraison()
                        n_entry.delete(0, 'end')
                        name_entry.delete(0, 'end')
                        adress_entry.delete(0, 'end')
                        prod_entry.delete(0, 'end') 
                    else: messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
                else: messagebox.showerror("Erreur", "L'ID de la livraison ne peut pas être modifié.")
                    



            def add_livraison():
            
                shipping_nbr = n_entry.get()
                client_name = name_entry.get()
                cin = cin_entry.get()
                shipping_address = adress_entry.get()
                phone_nbr = phone_nbr_entry.get()
                order_nbr = order_entry.get()
                order_date = order_date_entry.get_date() 
                order_date_l= order_date_l_entry.get_date() 
                central_dep_hr = central_dep_hr_combobox.get()
                central_arr_hr = central_arr_hr_combobox.get()
                worksite_dep_hr = worksite_dep_hr_combobox.get()
                worksite_arr_hr = worksite_arr_hr_combobox.get()    
                product_id = prod_entry.get()
                adjuvant=adjuvant_combobox.get()
                quantity = quantite_entry.get()
                quantity_a = quantity_a_entry.get()
                vehicle = matricule_entry.get()
                total_ttc = total_ttc_entry.get()
                chauffeur = chauffeur_combobox.get()
                variable1 = variable1_combobox.get() 
                if variable1 == "OUI":
                    variable1 ="X"
                else: variable1=" "
                variable2 = "X" if variable1 == " " else " " 

                if (shipping_nbr and client_name and cin and shipping_address and phone_nbr and order_nbr and order_date and central_dep_hr and central_arr_hr and
                    worksite_dep_hr and worksite_arr_hr and product_id and quantity and
                    vehicle and total_ttc and chauffeur and adjuvant and order_date_l 
                    and quantity_a and variable1):

                    if is_numeric_input(shipping_nbr):

                        data_string = f"{shipping_nbr},{client_name},{cin},{shipping_address},{phone_nbr}," \
                                    f"{order_nbr},{order_date},{central_dep_hr},{central_arr_hr}," \
                                    f"{worksite_arr_hr},{worksite_dep_hr},{product_id},{quantity},{adjuvant},{vehicle},{total_ttc},{variable1},{variable2}," \
                                    f"{chauffeur},{quantity_a},{order_date_l}"

                        # Create or open the CSV file in write mode
                        with open('livraison_data.csv', mode='a', newline='') as file:
                            writer = csv.writer(file, delimiter=';')

                            # Write the data to the CSV file
                            writer.writerow([data_string])
                        # Clear the entry fields after writing to CSV
                        root.withdraw()   
                        root.after(100, root.destroy)
                        n_entry.delete(0, 'end')
                        name_entry.delete(0, 'end')
                        adress_entry.delete(0, 'end')
                        prod_entry.delete(0, 'end') 
                        order_entry.delete(0, 'end')
                        display_livraison()
                    else: messagebox.showerror("Erreur", "Le numéro de livraison doit-être un chiffre !")
                else:
                    messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
                
            # Création de la fenêtre principale
            root = tk.Tk()
            root.title("Générateur de Bon de livraison")
            root.geometry('400x800')
            
            order_nbr = order_entry.get()
            client_name = name_entry.get()

            
            order_date_label = tk.Label(root, text="Date du bon de commande :")
            order_date_label.pack()
            order_date_entry = DateEntry(root, date_pattern="yyyy-mm-dd")
            order_date_entry.pack()

            order_date_l_label = tk.Label(root, text="Date du bon de livraison :")
            order_date_l_label.pack()
            order_date_l_entry = DateEntry(root, date_pattern="yyyy-mm-dd")
            order_date_l_entry.pack()

            adress_label = tk.Label(root, text="adresse :")
            adress_label.pack()
            adress_entry = tk.Entry(root)
            adress_entry.pack()

            phone_nbr_label = tk.Label(root, text="Phone Number:")
            phone_nbr_label.pack()
            phone_nbr_entry = tk.Entry(root)
            phone_nbr_entry.pack()

            cin_label = tk.Label(root, text="CIN:")
            cin_label.pack()
            cin_entry = tk.Entry(root)
            cin_entry.pack()  

            hours_label = tk.Label(root, text="Heure de départ centrale :")
            hours_label.pack()
            central_dep_hr_combobox = ttk.Combobox(root, values=["00:00"])  # Add your desired values
            central_dep_hr_combobox.pack()

            
            hours_label = tk.Label(root, text="Heure d'arrivée centrale :")
            hours_label.pack()
            central_arr_hr_combobox = ttk.Combobox(root, values=["00:00"])  # Add your desired values
            central_arr_hr_combobox.pack()

            hours_label = tk.Label(root, text="Heure de départ du site :")
            hours_label.pack()
            worksite_dep_hr_combobox = ttk.Combobox(root, values=["00:00"])  # Add your desired values
            worksite_dep_hr_combobox.pack()

            hours_label = tk.Label(root, text="Heure d'arrivée sur site :")
            hours_label.pack()
            worksite_arr_hr_combobox = ttk.Combobox(root, values=["00:00"])  # Add your desired values
            worksite_arr_hr_combobox.pack()
            
            # Zone de saisie pour la quantité
            quantite_label = tk.Label(root, text="Quantité :")
            quantite_label.pack()
            quantite_entry = tk.Entry(root)
            quantite_entry.pack()

            
            adjuvant_label = tk.Label(root, text="Adjuvant :")
            adjuvant_label.pack()
            adjuvant_combobox = ttk.Combobox(root, values=["Adjuvant 1", "Adjuvant 2", "Adjuvant 3"])  
            adjuvant_combobox.pack()

            quantity_a_label = tk.Label(root, text="Quantité d'adjuvant :")
            quantity_a_label.pack()
            quantity_a_entry = tk.Entry(root)
            quantity_a_entry.pack()

            variable1_label = tk.Label(root, text="Pompé :")
            variable1_label.pack()
            variable1_combobox = ttk.Combobox(root, values=["OUI", "NON"])  # Add your desired values
            variable1_combobox.pack()

        
            # Zone de saisie pour le camion
            matricule_label = tk.Label(root, text="Matricule :")
            matricule_label.pack()
            matricule_entry = ttk.Combobox(root, values=["Matricule 1", "Matricule 2", "Matricule 3"])
            matricule_entry.pack()
            
            chauffeur_label = tk.Label(root, text="Chauffeur :")
            chauffeur_label.pack()
            chauffeur_combobox = ttk.Combobox(root, values=["Chauffeur 1", "Chauffeur 2", "Chauffeur 3"])  
            chauffeur_combobox.pack()
            

            total_ttc_label = tk.Label(root, text="Total TTC:")
            total_ttc_label.pack()
            total_ttc_entry = tk.Entry(root)
            total_ttc_entry.pack()

            variable3_label = tk.Label(root, text="Signé :")
            variable3_label.pack()
            variable3_combobox = ttk.Combobox(root, values=["OUI", "NON"])  # Add your desired values
            variable3_combobox.pack()

            
            
            # Bouton pour générer le bon de livraison
            if R==2:
                selected_item = livraison_tree.selection()
                shp_nbr = livraison_tree.item(selected_item)["values"][0]
                generer_button = tk.Button(root, text="Modifier livraison",command=lambda : mod_livraison(shp_nbr))
                generer_button.pack()
            else:
                generer_button = tk.Button(root, text="Ajouter la livraison",command=add_livraison)
                generer_button.pack()
            
            # MODIfier livraison
            if R==2:
                selected_item = livraison_tree.selection()
                shp_nbr = livraison_tree.item(selected_item)["values"][0]
                mod_line=[]
                with open('livraison_data.csv', 'r', newline='') as file:
                        reader = csv.reader(file, delimiter=',')
                        lines = list(reader)
                for row in lines :
                    if str(row[0])==str(shp_nbr):
                        mod_line=row
                n_entry.delete(0, tk.END)
                n_entry.insert(tk.END, mod_line[0])

                name_entry.delete(0, tk.END)
                name_entry.insert(tk.END, mod_line[1])

                cin_entry.delete(0, tk.END)
                cin_entry.insert(tk.END, mod_line[2])

                adress_entry.delete(0, tk.END)
                adress_entry.insert(tk.END, mod_line[3])

                phone_nbr_entry.delete(0, tk.END)
                phone_nbr_entry.insert(tk.END, mod_line[4])

                order_entry.delete(0, tk.END)
                order_entry.insert(tk.END, mod_line[5])
        
                order_date_entry.set_date(mod_line[6])

                central_dep_hr_combobox.set(mod_line[7])

                central_arr_hr_combobox.set(mod_line[8])

                worksite_arr_hr_combobox.set(mod_line[10])

                worksite_dep_hr_combobox.set(mod_line[9])

                prod_entry.delete(0, tk.END)
                prod_entry.insert(tk.END, mod_line[11])

                adjuvant_combobox.set(mod_line[13])

                quantite_entry.delete(0, tk.END)
                quantite_entry.insert(tk.END, mod_line[15])

                quantity_a_entry.delete(0, tk.END)
                quantity_a_entry.insert(tk.END, mod_line[19])

                matricule_entry.delete(0, tk.END)
                matricule_entry.insert(tk.END, mod_line[14])

                total_ttc_entry.delete(0, tk.END)
                total_ttc_entry.insert(tk.END, mod_line[12])

                chauffeur_combobox.set(mod_line[18])
                if mod_line[16]=='X' : variable1_combobox.set('OUI')
                else : variable1_combobox.set('NON')
                order_date_l_entry.set_date(mod_line[20])
            
            else:
                if order_nbr:
                    with open('class_order.csv', 'r', newline='') as file:
                        reader = csv.reader(file, delimiter=',')
                        lines = list(reader)
                    for row in lines :
                        if str(row[0])==str(order_nbr):
                            mod_line=row
                    original_date_str = mod_line[1].replace('/', '-')
                    
                    parts = original_date_str.split("-")

                    formatted_date = parts[2] + "-" + parts[1] + "-" + parts[0]

                    order_date_entry.set_date(formatted_date)

                if client_name:
                    with open('class_client.csv', 'r', newline='') as file:
                        reader = csv.reader(file, delimiter=',')
                        lines = list(reader)
                    for row in lines :
                        if str(row[1])==str(client_name):
                            mod_line=row
                    phone_nbr_entry.insert(tk.END, mod_line[3])
                    adress_entry.insert(tk.END, mod_line[2])

            # Lancement de l'application
            root.mainloop()

        def bon_livraison():
          
            selected_item = livraison_tree.selection()
            shp_nbr = livraison_tree.item(selected_item)["values"][0]
            livraison_instances = []
            if shp_nbr:
                with open('livraison_data.csv', 'r') as file:
                    reader = csv.reader(file, delimiter=',')
                    # Skip the header row
                    next(reader)
                    for row in reader:
                        # Instantiate a Livraison object for each row
                        livraison_instance = livraison(*row)
                        livraison_instances.append(livraison_instance)

                    for ligne in livraison_instances:
                    
                            if int(shp_nbr)== int(ligne.shipping_nbr()):

                                remplir_template(ligne)
                
                                convert_docx_to_pdf(str(ligne.shipping_nbr())+'.docx', str(ligne.shipping_nbr())+'.pdf')
                            
        
        def display_livraison():
    # Clear existing data in the tree
            livraison_tree.delete(*livraison_tree.get_children())
            
            # Read data from the CSV file
            with open('livraison_data.csv', 'r') as file:
            
                reader = csv.reader(file, delimiter=',')
                next(reader)
                for livraison in reader:
                    # Display specific columns (adjust the column indices based on your CSV structure)
                    livraison_tree.insert("", tk.END, values=(livraison[0], livraison[1], livraison[3], livraison[11]))

       
        entry_var = tk.StringVar()
        titre_label = tk.Label(frame, text="Livraison", font=("Arial", 16))
        titre_label.pack(pady=5)
        
        columns_livraison = ("N° de livraison", "Client", "Adresse", "ID produit")
        
        tree_frame = tk.Frame(frame)
        tree_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical")
        scrollbar.pack(side='right', fill='y')

        style = ttk.Style()
        style.configure('Treeview', rowheight=25)

        
        livraison_tree = ttk.Treeview(tree_frame, columns=columns_livraison, show="headings", style='Custom.Treeview')

        scrollbar.config(command=livraison_tree.yview)

        
        for col in columns_livraison:
            livraison_tree.heading(col, text=col)
            livraison_tree.column(col, width=150)
        
        livraison_tree.pack(fill=tk.BOTH, expand=True, pady=10)
        livraison_tree.bind("<Double-1>", double_click_livraison)

        class AutocompleteEntry(ttk.Combobox):
            def set_completion_list(self, completion_list):
                self._completion_list = sorted(completion_list)
                self._hits = []
                self.position = 0
                self.bind('<KeyRelease>', self.handle_keyrelease)
                self['values'] = self._completion_list

            def autocomplete(self, delta=0):
                if delta:
                    self.delete(self.position, tk.END)
                else:
                    self.position = len(self.get())

                _hits = []
                for item in self._completion_list:
                    if item.lower().startswith(self.get().lower()):
                        _hits.append(item)

                if _hits != self._hits:
                    self._hits = _hits
                    self['values'] = _hits

            def handle_keyrelease(self, event):
                if event.keysym in ('BackSpace', 'Left', 'Right', 'Up', 'Down', 'Shift', 'Control'):
                    return

                if event.keysym == 'Return':
                    self._hits = []
                    return

                if event.keysym in ('BackSpace', 'Left', 'Right'):
                    self.autocomplete(-1)
                else:
                    self.autocomplete()

        def load_names_from_csv(csv_file):
            names = []
            with open(csv_file, 'r') as file:
                reader = csv.reader(file, delimiter=',')
                next(reader)  # Skip header
                for row in reader:
                    names.append(row[1])  # Assuming names are in the first column
            return names

        def on_select(event):
            value = entry_var.get()
            print(f"Selected: {value}")

        names_list=load_names_from_csv('class_client.csv')

        # Add the labels and input fields for adding/modifying a client
        input_frame_livraison = tk.Frame(frame)
        input_frame_livraison.pack()
        
        n_label = tk.Label(input_frame_livraison, text="N° de livraison :")
        n_label.pack(side=tk.LEFT, padx=5)
        n_entry = tk.Entry(input_frame_livraison)
        n_entry.pack(side=tk.LEFT, padx=5)
        
        
        name_label = tk.Label(input_frame_livraison, text="Nom :")
        name_label.pack(side=tk.LEFT, padx=5)
        name_entry = AutocompleteEntry(input_frame_livraison, textvariable=entry_var)
        name_entry.set_completion_list(names_list)
        name_entry.pack(side=tk.LEFT, padx=5)
        

        
        order_label = tk.Label(input_frame_livraison, text="N° de commande :")
        order_label.pack(side=tk.LEFT, padx=5)
        order_entry = tk.Entry(input_frame_livraison)
        order_entry.pack(side=tk.LEFT, padx=5)

        prod_label = tk.Label(input_frame_livraison, text="ID Produit :")
        prod_label.pack(side=tk.LEFT, padx=5)
        prod_entry = tk.Entry(input_frame_livraison)
        prod_entry.pack(side=tk.LEFT, padx=5)
        
        # Add the buttons for adding/modifying a client
        button_frame_livraison = tk.Frame(frame)
        button_frame_livraison.pack(pady=10)

        display_liv = ctk.CTkButton(button_frame_livraison, text="Liste des livraisons", command=display_livraison)
        display_liv.pack(side=tk.LEFT, padx=5)
        
        add_button_livraison = tk.Button(button_frame_livraison, text="Ajouter Livraison", command=lambda:page_livraison(1))
        add_button_livraison.pack(side=tk.LEFT, padx=5)
        
        delete_button_livraison = tk.Button(button_frame_livraison, text="Supprimer Livraison", command= supress_livraison)
        delete_button_livraison.pack(side=tk.LEFT, padx=5)
        
        modify_button_livraison = tk.Button(button_frame_livraison, text="Modifier Livraison",command= lambda: page_livraison(2))
        modify_button_livraison.pack(side=tk.LEFT, padx=5)

        modify_button_livraison = tk.Button(button_frame_livraison, text="Créer bon de Livraison", command=bon_livraison)
        modify_button_livraison.pack(side=tk.LEFT, padx=5)
        
    def Clients():
        Clear_widgets(frame)

        def generate_new_client_id():
            # Retrieve existing supplier IDs from the CSV file
            deleted_clients_file_path = "./deleted_clients.csv"
            client_file_path = "./class_client.csv"  # Replace with the actual path to your CSV file
            existing_ids = set()
            deleted_ids = set()

            # reading the current ids
            with open(client_file_path, newline='', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile)
                for client in reader:
                    existing_ids.add(int(client['ID']))
            
            # reading the deleted ids
            try:
                with open(deleted_clients_file_path, newline='', encoding='utf-8') as csvfile:
                    reader = csv.DictReader(csvfile)
                    for client in reader:
                        deleted_ids.add(int(client['ID'])) 
            except FileNotFoundError:
                # if there is no deleted clients file, just ignore
                pass

            all_ids = existing_ids.union(deleted_ids)
            new_id = max(all_ids, default=0) + 1
            return new_id

        def display_clients():
            clients_tree.delete(*clients_tree.get_children())
            if not clients_tree.get_children():
                for client in clients_data:
                    clients_tree.insert("", tk.END, values=(client["ID"], client["Name"], client["Address"], client["Email"], client["Phone Number"]))
        
        def refresh_clients_table():
            clients_tree.delete(*clients_tree.get_children())
            for client in clients_data:
                clients_tree.insert("", tk.END, values=(client["ID"], client["Name"], client["Address"], client["Email"], client["Phone Number"]))
        
        
        def double_click_client(event):
            selected_item = clients_tree.focus()
            if selected_item:
                values = clients_tree.item(selected_item, "values")
                if values:
                    id_entry.delete(0, tk.END)
                    client_name_entry.delete(0, tk.END)
                    client_address_entry.delete(0,tk.END)
                    email_entry.delete(0, tk.END)
                    phone_entry.delete(0, tk.END)
                    id_entry.insert(tk.END, values[0])
                    client_name_entry.insert(tk.END, values[1])
                    client_address_entry.insert(tk.END, values[2])
                    email_entry.insert(tk.END, values[3])
                    phone_entry.insert(tk.END, values[4])
            
        def delete_client():
            selected_item = clients_tree.selection()
            if selected_item:
                item_id = clients_tree.item(selected_item)["values"][0]
                df = pd.read_csv('./class_client.csv', sep = ',')

                deleted_client = df[df['ID'] == item_id]
                # date of suppression
                deleted_client['DateSuppr'] = date.today().strftime("%d/%m/%Y")
                # add the deleted client to the deleted_clients csv file for archive
                with open('./deleted_clients.csv', 'a', newline='') as f:
                    deleted_client.to_csv(f, header=f.tell()==0, index=False)


                df = df[df['ID'] != item_id]
                df.to_csv('./class_client.csv', index=False)
                for client in clients_data:
                    if client["ID"] == item_id:
                        clients_data.remove(client)
                        break
                
                clients_tree.delete(selected_item)
                refresh_client_ids()
                refresh_clients_table()
            else:
                messagebox.showwarning("Avertissement", "Veuillez sélectionner un client à supprimer.")

        
        def add_client():
            name = client_name_entry.get()
            email = email_entry.get()
            phone_entry_value = phone_entry.get()
            address = client_address_entry.get()

            if len(phone_entry_value) == 11:  # Vérifie si la longueur est de 11 chiffres
                phone_number = phone_entry_value  # Utilise la valeur telle quelle
            else:
                phone_number = '216' + phone_entry_value  # Ajoute '216' au début

            if not is_valid_phone_number(phone_number):
                    messagebox.showerror("Erreur", "Le numéro de téléphone ne doit pas commencer par zéro et doit contenir 11 chiffres.")
                    return  # Exit the function if the phone number is not valid
   
            if not is_valid_email(email):
                    messagebox.showerror("Erreur", "L'adresse e-mail n'est pas valide. \n Format : ___@___.__")
                    return  # Exit the function if the email address is not valid
            
            



            if name and email and phone_number:
                if is_numeric_input(phone_number):
                    df = pd.read_csv('./class_client.csv')
                    size = df.shape[0] + 1

                    client_id= generate_new_client_id()
                    new_client = Client(client_id, name, address, email, phone_number)

                    df.loc[size] = [client_id, name, address, email, phone_number]
                    df.to_csv('./class_client.csv', index = False)
           
           
                    client_instances.append(new_client)
                    clients_tree.insert("", tk.END, values=(client_id, name, address, email, phone_number))
                    client_name_entry.delete(0, tk.END)
                    client_address_entry.delete(0,tk.END)
                    email_entry.delete(0, tk.END)
                    phone_entry.delete(0, tk.END)
                else:
                    messagebox.showerror("Erreur", "Le numéro de téléphone doit être une valeur numérique.")
            else:
                messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
                messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")

        def modify_client():
            selected_item = clients_tree.selection()
            if selected_item:
                client_id = id_entry.get()
                name = client_name_entry.get()
                address = client_address_entry.get()
                email = email_entry.get()
                phone_number = phone_entry.get()

                if client_id and name and email and phone_number:
                    if is_numeric_input(client_id) and is_numeric_input(phone_number):
                        client_id = int(client_id)
                        selected_client_id = clients_tree.item(selected_item)["values"][0]
                        if selected_client_id == client_id:
                    
                            df = pd.read_csv('./class_client.csv')
                            colonne_index = 'ID'
                            df = df.set_index(colonne_index)
                            nouvelles_valeurs = {'Name': name, 'Address': address, 'Email': email, 'Phone Number': phone_number}
                            df.loc[client_id] = nouvelles_valeurs
                            df.reset_index(inplace = True)
                    
                            df.to_csv('./class_client.csv', index = False)
                    
                            clients_tree.item(selected_item, values=(client_id, name, address, email, phone_number))
                            messagebox.showinfo("Succès", "Client modifié avec succès.")
                    
                        else:
                            messagebox.showerror("Erreur", "L'ID du client ne peut pas être modifié.")
                    else:
                        messagebox.showerror("Erreur", "Le numéro de téléphone et l'ID doivent être des valeurs numériques.")
                else:
                    messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
            else:
                messagebox.showwarning("Avertissement", "Veuillez sélectionner un client à modifier.")

        def search_client(search_value, attribute):
            # search_value = search_entry.get()

            # with open('class_client.csv', mode='r', encoding='utf-8') as file:
            #     # Créer un objet reader pour lire le fichier CSV
            #     csv_reader = csv.DictReader(file)
        
            #     # Initialiser une variable pour vérifier si on a trouvé au moins un client correspondant
            #     found = False

            #     # Initialiser une liste pour stocker les clients trouvés
            #     found_clients = []

            #     # Parcourir chaque ligne du fichier CSV
            #     for row in csv_reader:
            #         # Vérifier si la valeur recherchée correspond à l'attribut spécifié (Nom ou ID)
            #         if search_value.lower() in row[attribute].lower():
            #             found_clients.append(row)
        
            #             # Si aucun client correspondant n'a été trouvé, afficher un message
            #     if not found:
            #         print("Aucun client correspondant à la recherche.")
            # Supprimez les anciens résultats de recherche
            clients_tree.delete(*clients_tree.get_children())

            # Dictionnaire pour stocker les résultats de recherche
            search_results = []
    
            # Recherche insensible à la casse
            search_value_lower = search_value.lower()

            # Parcourir le dictionnaire clients_data à la recherche de correspondances
            for client in clients_data:
                if search_value_lower in str(client[attribute]).lower():
                    search_results.append(client)
    
            # Affichage des résultats de recherche
            if search_results:
                for client in search_results:
                    clients_tree.insert("", tk.END, values=(client["ID"], client["Name"], client["Address"], client["Email"], client["Phone Number"]))
            else:
                messagebox.showinfo("Résultat de recherche", "Aucun client correspondant à la recherche.")

        def search_clients_wrapper():
            search_value = search_entry.get()
            search_attribute_value = search_attribute.get()
            search_client(search_value, search_attribute_value)

        


        def open_deleted_clients_window():

            def display_deleted_clients():
                deleted_clients_tree.delete(*deleted_clients_tree.get_children())
                if not deleted_clients_tree.get_children():
                    for client in deleted_clients_data:
                        deleted_clients_tree.insert("", tk.END, values=(client["ID"], client["Name"], client["Address"], client["Email"], client["Phone Number"],client['DateSuppr']))
        
            # Crée une nouvelle fenêtre pour afficher les paramètres
            new_window = tk.Tk()
            new_window.title("Clients Supprimés")

            # tk.Label(new_root, text="Entrez la durée en jours:").pack()

            # duration_entry = tk.Entry(reset_window)
            # duration_entry.pack()

            # confirm_button = tk.Button(reset_window, text="Confirmer", command=lambda: reset_history(duration_entry.get()))
            # confirm_button.pack()

            deleted_clients_frame = tk.Frame(new_window)
            deleted_clients_frame.pack()

            deleted_clients_tree_frame = tk.Frame(deleted_clients_frame)
            deleted_clients_tree_frame.pack(fill=tk.BOTH, expand=True, pady=10)

            

            # Création du Treeview dans la fenêtre popup
            columns_deleted_clients = ("ID", "Nom", "Adresse", "Email", "Phone Number", "DeletionDate")
            deleted_clients_tree = ttk.Treeview(deleted_clients_tree_frame, columns=columns_deleted_clients, show="headings")
    
            # Configuration des colonnes
            for col in columns_deleted_clients:
                deleted_clients_tree.heading(col, text=col)
                deleted_clients_tree.column(col, width=120)

            # Ajout d'un scrollbar
            scrollbar = ttk.Scrollbar(deleted_clients_tree_frame, orient="vertical", command=deleted_clients_tree.yview)
            deleted_clients_tree.configure(yscroll=scrollbar.set)
            scrollbar.pack(side='right', fill='y')
            deleted_clients_tree.pack(fill='both', expand=True)

            # Remplissez le Treeview avec les données des clients supprimés
            # ... (Votre logique pour remplir le Treeview)

            # Ajout des boutons de gestion du délai de suppression
            # buttons
            button_frame_deleted_clients = tk.Frame(deleted_clients_frame)
            button_frame_deleted_clients.pack()

            display_deleted_clients = ctk.CTkButton(button_frame_deleted_clients, text="Liste des Clients Supprimés", command=display_deleted_clients)
            display_deleted_clients.pack(side=tk.LEFT, padx=5)
            # delay_frame = tk.Frame(new_root)
            # delay_frame.pack(pady=5)

            delay_label = tk.Label(button_frame_deleted_clients, text="Délai de suppression (nb semaines):")
            delay_label.pack(side=tk.LEFT, padx=5)

            delay_entry = tk.Entry(button_frame_deleted_clients)
            delay_entry.pack(side=tk.LEFT, padx=5)

            set_delay_button = tk.Button(button_frame_deleted_clients, text="Définir Délai", command=lambda: set_deletion_delay())
            set_delay_button.pack(side=tk.LEFT, padx=5)

            # Définissez la fonction pour gérer le délai de suppression ici
            def set_deletion_delay():
                try:
                    deletion_delay = int(delay_entry.get())
                    # Enregistrer le délai en semaines dans le fichier
                    with open('deletion_delay_weeks.txt', 'w') as f:
                        f.write(str(deletion_delay))
                    messagebox.showinfo("Succès", f"Délai de suppression défini à {deletion_delay} semaines.")
                except ValueError:
                    messagebox.showerror("Erreur", "Veuillez entrer un nombre entier valide.")

            

            # Assurez-vous que la fenêtre popup est modale
            new_window.grab_set()
            new_window.focus_set()
            new_window.wait_window()

        
        def reset_history(duration):
            try:
                duration_in_days = int(duration)
                # Ici, vous implémenterez la logique pour nettoyer l'historique basé sur la durée
                # Par exemple, vous pouvez parcourir le fichier CSV et supprimer les entrées plus anciennes que la durée spécifiée
                print(f"Réinitialisation de l'historique des clients supprimés pour les enregistrements plus vieux que {duration_in_days} jours.")
            except ValueError:
                messagebox.showerror("Erreur", "Veuillez entrer un nombre valide.")


        


        

        
        # def total_spent(start_date=None, end_date=None):
        #     selected_client = clients_tree.selection()
        #     if selected_client:
        #         client_id = clients_tree.item(selected_client)["values"][0] 

        
        #         client_sales = [sale for sale in sales_data if sale["Client ID"] == client_id]
        #         # Si des dates de début et de fin sont spécifiées, filtrer les ventes dans cette plage (à implémenter)
        #         if start_date and end_date:
        #             start_date = datetime.datetime.strptime(start_date, "%Y-%m-%d")
        #             end_date = datetime.datetime.strptime(end_date, "%Y-%m-%d")
        #             client_sales = [sale for sale in client_sales if start_date <= datetime.datetime.strptime(sale["Sale Date"], "%Y-%m-%d") <= end_date]

        #         # Calculate the total amount spent
        #         total_amount = sum(float(sale["Sale Price"]) * int(sale["Quantity Sold"]) for sale in client_sales)

        #         # Display 
        #         messagebox.showinfo("Total dépensé", f"Le client {client_id} a dépensé un total de {total_amount:.2f} unités.")
        #     else:
        #         messagebox.showwarning("Avertissement", "Veuillez sélectionner un client.")



        titre_label = tk.Label(frame, text="Clients", font=("Arial", 16))
        titre_label.pack(pady=5)
        
        columns_clients = ("ID", "Nom", "Address", "Email", "Numéro de téléphone")
        
        tree_frame = tk.Frame(frame)
        tree_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical")
        scrollbar.pack(side='right', fill='y')
        
        style = ttk.Style()
        style.configure('Treeview', rowheight=25)
        global clients_tree
        clients_tree = ttk.Treeview(tree_frame, columns=columns_clients, show="headings", style='Custom.Treeview')
        
        scrollbar.config(command=clients_tree.yview)    
        clients_tree = ttk.Treeview(tree_frame, columns=columns_clients, show="headings", style='Custom.Treeview')
        scrollbar.config(command=clients_tree.yview)
        # clients_tree.pack(side='left', fill=tk.BOTH, expand=True)

        for col in columns_clients:
            clients_tree.heading(col, text=col)
            clients_tree.column(col, width=150)
    
        clients_tree.pack(fill=tk.BOTH, expand=True, pady=5)
        clients_tree.bind("<Double-1>", double_click_client)

        input_frame_clients = tk.Frame(frame)
        input_frame_clients.pack()

        client_id_label = tk.Label(input_frame_clients, text="ID de client :")
        client_id_label.pack(side=tk.LEFT, padx=5)
        id_entry = tk.Entry(input_frame_clients, validate="key")
        id_entry.config(validatecommand=(frame.register(validate_id), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
        id_entry.pack(side=tk.LEFT, padx=5)

        name_label = tk.Label(input_frame_clients, text="Nom :")
        name_label.pack(side=tk.LEFT, padx=5)
        client_name_entry = tk.Entry(input_frame_clients)
        client_name_entry.pack(side=tk.LEFT, padx=5)

        client_address_label = tk.Label(input_frame_clients, text="Adresse :")
        client_address_label.pack(side=tk.LEFT, padx=5)
        client_address_entry = tk.Entry(input_frame_clients)
        client_address_entry.pack(side=tk.LEFT, padx=5)
    
        email_label = tk.Label(input_frame_clients, text="E-mail :")
        email_label.pack(side=tk.LEFT, padx=5)
        email_entry = tk.Entry(input_frame_clients)
        email_entry.pack(side=tk.LEFT, padx=5)
    
        phone_label = tk.Label(input_frame_clients, text="Numéro de téléphone :")
        phone_label.pack(side=tk.LEFT, padx=5)
        phone_entry = tk.Entry(input_frame_clients, validate="key")
        phone_entry.config(validatecommand=(frame.register(validate_phone_number), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
        phone_entry.pack(side=tk.LEFT, padx=5)

        search_frame = tk.Frame(frame)
        search_frame.pack(pady=5)


        search_label = tk.Label(search_frame, text="Rechercher")
        search_label.pack(side=tk.LEFT, padx=5)
        search_entry = tk.Entry(search_frame)
        search_entry.pack(side=tk.LEFT, padx=5)


        # buttons
        button_frame_clients = tk.Frame(frame)
        button_frame_clients.pack(pady=10)

        display_clients = ctk.CTkButton(button_frame_clients, text="Liste des Clients", command=display_clients)
        display_clients.pack(side=tk.LEFT, padx=5)
    
        add_button_clients = tk.Button(button_frame_clients, text="Ajouter Client", command=add_client)
        add_button_clients.pack(side=tk.LEFT, padx=5)
    
        delete_button_clients = tk.Button(button_frame_clients, text="Supprimer Client", command=delete_client)
        delete_button_clients.pack(side=tk.LEFT, padx=5)
    
        modify_button_clients = tk.Button(button_frame_clients, text="Modifier Client", command=modify_client)
        modify_button_clients.pack(side=tk.LEFT, padx=5)

        deleted_clients_button = tk.Button(button_frame_clients, text="Clients Supprimés", command=open_deleted_clients_window)
        deleted_clients_button.pack(side=tk.LEFT, padx=5)

        search_button = tk.Button(search_frame, text="Rechercher", command=search_clients_wrapper)
        search_button.pack(side=tk.LEFT, padx=5)

        # Créer un menu déroulant pour choisir l'attribut de recherche
        search_attribute = tk.StringVar()
        search_attribute.set("Name")  # Valeur par défaut
        search_options = ttk.Combobox(search_frame, textvariable=search_attribute)
        search_options['values'] = ("ID", "Name")
        search_options.pack()

        exit_button = tk.Button(button_frame_clients, text="Quitter", command=frame.quit)
        exit_button.pack(side=tk.LEFT, padx=5)


    #//////////////////////////// INTERFACE ORDER ///////////////////////////////////////////
    def orders():    
        Clear_widgets(frame)

        def add_order():
            client_id = order_client_id_entry.get()
            if client_id.isnumeric():
                client = getClientById(int(client_id))
                if client:
                    order_date = order_date_label.cget('text')
                    type_transaction = type_transaction_var.get()
                    statut = statut_var.get()
                    pompe = True if pompe_var.get() == 'Oui' else False

                    if order_date != "" and order_listprod_var.get()!="":
                        products = [(getProductByDescription(desc), qty) for desc, qty in (a.split(' : ') for a in order_listprod_var.get().split('\n'))]
                        pricePaid = float(order_paid_entry.get()) if is_float(order_paid_entry.get()) else 0.0
                        new_id = get_next_order_id()
                        new_order = Order(new_id, order_date, client, products, type_transaction, pompe=pompe, statut=statut, price_paid=pricePaid)
                        if new_order.price==new_order.price_paid:
                            new_order.statut="Payée"
                        order_instances.append(new_order)
                        orders_tree.insert("", tk.END, values=(new_id, order_date, client.id, new_order.get_str_Products(), type_transaction, 'Oui' if pompe else 'Non', statut, new_order.price_paid, new_order.price))
        
                        df = pd.read_csv('./class_order.csv')
                        size = df.shape[0] + 1
                        df.loc[size] = [new_id, order_date, client_id, new_order.get_str_Products(id=True), payment_type, order.price, order.price_paid, pompe, statut]
                        df.to_csv('./class_order.csv', index=False)

                        resetVar()
                    else:
                        messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
                else:
                    # Le client n'existe pas, afficher un message d'erreur
                    messagebox.showerror("Client inexistant", "Le client n'existe pas. Veuillez le créer avant d'ajouter la commande.")
            else:
                messagebox.showerror("Erreur", "L'identifiant de client doit être une valeur numérique.")

        def modify_order():
            selected_item = orders_tree.selection()
            if selected_item:
                order_id = order_id_entry.get()
                order_date = order_date_label.cget('text')
                client_id = order_client_id_entry.get()
                type_transaction = type_transaction_var.get()
                statut = statut_var.get()
                pompe = True if pompe_var.get() == 'Oui' else False
                if order_id and order_date and client_id and order_listprod_var.get()!="" and type_transaction and statut:
                    if is_numeric_input(order_id) and is_numeric_input(client_id):
                        order_id, client_id = int(order_id), int(client_id)
                        products = [(getProductByDescription(desc), qty) for desc, qty in (a.split(' : ') for a in order_listprod_var.get().split('\n'))]
                        pricePaid = float(order_paid_entry.get()) if is_float(order_paid_entry.get()) else 0.0
                        index = getOrderById(order_id, index=True)
                        client = getClientById(client_id)
                        selected_order_id = orders_tree.item(selected_item)["values"][0]
                        if selected_order_id == order_id:
                            if index!=None and client!=None:
                                order = Order(order_id, order_date, client, products, type_transaction, pompe=pompe, statut=statut, price_paid=pricePaid)
                                if order.price==order.price_paid:
                                    order.statut="Payée"
                                order_instances[index] = order
                                df = pd.read_csv('./class_order.csv')
                                df.set_index('order_id', inplace=True)
                                df.loc[order_id] = {'order_date': order_date, 'client_id': client_id, 'products' : order.get_str_Products(id=True),
                                                'payment_type': type_transaction,'price': order.price,'price_paid': order.price_paid,'pompe': pompe,'statut': order.statut}
                                df.reset_index(inplace = True)
                                df.to_csv('./class_order.csv', index = False)

                                orders_tree.item(selected_item, values=(order_id, order_date, client_id, order.get_str_Products(), type_transaction, 'Oui' if pompe else 'Non', order.statut, order.price_paid, order.price))
                                
                                resetVar()
                                messagebox.showinfo("Succès", "Commande modifiée avec succès.")
                            else:
                                messagebox.showerror("Erreur", "ID de commande ou de client inexistant")
                        else:
                            messagebox.showerror("Erreur", "L'ID de la commande ne peut pas être modifié.")
                    else:
                        messagebox.showerror("Erreur", "L'identifiant de commande et l'identifiant de client doivent être des valeurs numériques.")
                else:
                    messagebox.showerror("Erreur", "Veuillez remplir tous les champs !")
            else:
                messagebox.showwarning("Avertissement", "Veuillez sélectionner une commande à modifier.")
        
        def delete_order():
            selected_item = orders_tree.selection()
            if selected_item:
                order_id = orders_tree.item(selected_item)["values"][0]
                order=getOrderById(order_id)
                if order!=None:
                    order_instances.remove(order)
                    df = pd.read_csv('./class_order.csv', sep = ',')
                    df = df[df['order_id'] != order_id]
                    df.to_csv('./class_order.csv', index=False)
                orders_tree.delete(selected_item)
            else:
                messagebox.showwarning("Avertissement", "Veuillez sélectionner une commande à supprimer.")
        
        def double_click_order(event):
            selected_item = orders_tree.focus()
            if selected_item:
                values = orders_tree.item(selected_item, "values")
                if values:
                    order_id_entry.delete(0, tk.END)
                    order_client_id_entry.delete(0, tk.END)
                    order_id_entry.insert(tk.END, values[0])
                    order_date_label.config(text=values[1])
                    order_client_id_entry.insert(tk.END, values[2])
                    order_listprod_var.set("\n".join(values[3].split(" | ")))
                    type_transaction_var.set(values[4])
                    pompe_var.set(values[5])
                    statut_var.set(values[6])
                    order_paid_entry.delete(0, tk.END)
                    order_paid_entry.insert(0, values[7])
        
        def resetVar():
            order_id_entry.delete(0, tk.END)
            order_date_label.config(text="")
            order_client_id_entry.delete(0, tk.END)
            order_listprod_var.set("")
            type_transaction_var.set("Chèque")
            pompe_var.set("Oui")
            statut_var.set("Non Payée")
            order_paid_entry.delete(0, tk.END)
        
        def GetDate(event): #create a new window with a calendar
            newWindow = tk.Toplevel(frame)
            newWindow.geometry('250x200+300+100')
            today = date.today()
            cal = Calendar(newWindow, selectmode = 'day', year = today.year, month = today.month, day = today.day, date_pattern='dd/mm/y')
            cal.pack(pady=5)
            
            def onclick_date(event):
                order_date_label.config(text=cal.get_date())
                newWindow.destroy()
            
            for row in cal._calendar:
                for lbl in row:
                    lbl.bind('<Double-1>', onclick_date)
        
        def ProductList(event): #create a new window to headle the product/qty list
            def onclickProduct(event=None):
                if order_add_qty_entry.get():
                    product, qty = order_add_product.get(), order_add_qty_entry.get()
                    changed=False
                    for i in range(len(listProd)):
                        if listProd[i][0]==product:
                            listProd[i] = (product, qty)
                            changed=True
                            break
                    if changed:
                        order_listprod_var.set("\n".join(" : ".join(e for e in tup) for tup in listProd))
                    else:
                        s = order_listprod_var.get()+'\n' if order_listprod_var.get()!="" else ""
                        s+=f"{product} : {qty}"
                        listProd.append((product, qty))
                        order_listprod_var.set(s)
                    order_add_qty_entry.delete(0, tk.END)
                    order_add_qty_entry.focus_set()
            
            def Suppr():
                order_listprod_var.set("")
                listProd=[]
            
            newWindow = tk.Toplevel(frame)
            newWindow.geometry('380x150+400+200')
            tk.Label(newWindow, text='ID Produit : Quantité (kg/m\u00B3)').grid(column=0, row=0) #^3 : U+00B3
            tk.Label(newWindow, textvariable=order_listprod_var, bg='white', bd=1, justify='left', anchor='w', relief='sunken').grid(column=0, row=1, columnspan=2, sticky='WE', pady=3)
            order_add_product = tk.StringVar(newWindow, product_instances[0].description)
            ttk.Combobox(newWindow, textvariable= order_add_product, values=[prod.description for prod in product_instances], state="readonly").grid(column=0, row=2)
            order_add_qty_entry = tk.Entry(newWindow, width=14, validate='key', validatecommand=(frame.register(validate_id), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
            order_add_qty_entry.grid(column=1, row=2)
            order_add_qty_entry.bind('<Return>', onclickProduct)
            tk.Button(newWindow, text="Supprimer", command=Suppr).grid(column=0, row=3)
            tk.Button(newWindow, text="Ajouter produit", command=onclickProduct).grid(column=1, row=3, pady=5)
            listProd = [(desc,qty) for desc, qty in (a.split(' : ') for a in order_listprod_var.get().split('\n'))] if order_listprod_var.get()!="" else []

        def display_orders():
            orders_tree.delete(*orders_tree.get_children())
            for order in order_instances:
                orders_tree.insert("", tk.END, values=(
                    order.order_id, order.order_date, order.client.id, order.get_str_Products(),
                    order.payment_type, 'Oui' if order.pompe else 'Non', order.statut, order.price_paid, order.price))
        
        # historique par client
        def get_order_history():
            client_id = int(order_client_id_entry.get()) if order_client_id_entry.get().isnumeric() else None
            if client_id:
                client_orders = [order for order in order_instances if order.client.clt_id == client_id]
                if len(client_orders)!=0:
                    orders_tree.delete(*orders_tree.get_children())
                    for order in client_orders:
                        orders_tree.insert("", tk.END, values=(
                            order.order_id, order.order_date, order.client.clt_id, order.get_str_Products(),
                            order.payment_type, 'Oui' if order.pompe else 'Non', order.statut, order.price_paid, order.price))
                else:
                    messagebox.showinfo("Information", f"Aucune commande pour ID : {client_id}")
                order_client_id_entry.delete(0, tk.END)
            else:
                messagebox.showerror("Erreur", "ID client incorrect")

        def get_unpaid_orders():
            show=True
            client_id = int(order_client_id_entry.get()) if order_client_id_entry.get().isnumeric() else None
            if client_id: #unpaid orders by client_id
                if any(c.clt_id==client_id for c in client_instances):
                    unpaid_orders = [order for order in order_instances if order.client.clt_id==client_id and order.statut=="Non Payée"]
                else:
                    messagebox.showerror("Erreur", 'ID client incorrect')
                    show=False
                order_client_id_entry.delete(0, tk.END)
            else: #all unpaid orders
                unpaid_orders = [order for order in order_instances if order.statut=="Non Payée"]
            if show:
                if len(unpaid_orders)!=0:
                    orders_tree.delete(*orders_tree.get_children())
                    for order in unpaid_orders:
                        orders_tree.insert("", tk.END, values=(
                            order.order_id, order.order_date, order.client.clt_id, order.get_str_Products(),
                            order.payment_type, 'Oui' if order.pompe else 'Non', order.statut, order.price_paid, order.price))
                        
                elif client_id:
                    messagebox.showinfo("Information", "Ce client n'a pas de commandes impayées")
                else:
                    messagebox.showinfo("Information", "Aucune commande impayée")
                
        def buttonBDC():
            orderid = order_id_entry.get()
            if is_numeric_input(orderid):
                try:
                    order = getOrderById(int(orderid))
                    try:
                        order.CreateBDC()
                        try:
                            convert_docx_to_pdf(f'./Bon de Commandes/Bon de Commande - {order.order_id}.docx', f'./Bon de Commandes/Bon de Commande - {order.order_id}.pdf')
                            os.remove(f'./Bon de Commandes/Bon de Commande - {order.order_id}.docx')
                        except:
                            messagebox.showerror("Erreur", "Erreur conversion word en pdf")
                    except:
                        messagebox.showerror("Erreur", "Erreur création Bon de Commande")
                except:
                    messagebox.showerror("Erreur", "ID commande incorrect")
                
            else:
                messagebox.showwarning("Avertissement", "Veuillez sélectionner un client à modifier.")

        titre_label = tk.Label(frame, text="Orders", font=("Arial", 16))
        titre_label.pack(pady=5)
        tree_frame = tk.Frame(frame)
        tree_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        columns = ("commande ID", "Date de la commande", "Client ID", "Produits", "Type de Transaction", "Pompé", "Status", "Montant payé", "Montant")
        order_col_size = [90, 130, 60, 220, 120, 80, 80, 100, 100]
        
        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical")
        scrollbar.pack(side='right', fill='y')
        
        global orders_tree
        orders_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", yscrollcommand=scrollbar.set)
        scrollbar.config(command=orders_tree.yview)
        
        for i,col in enumerate(columns):
            orders_tree.heading(col, text=col)
            orders_tree.column(col, width=order_col_size[i])

        orders_tree.pack(side=tk.LEFT, pady=10)
        orders_tree.bind("<Double-1>", double_click_order)
        
        
        # Add the labels and input fields for adding/modifying an order
        input_frame_orders = tk.Frame(frame)
        input_frame_orders.pack()

        tk.Label(input_frame_orders, text="ID de commande :").pack(side=tk.LEFT, padx=5)
        order_id_entry = tk.Entry(input_frame_orders, validate="key")
        order_id_entry.config(validatecommand=(frame.register(validate_id), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
        order_id_entry.pack(side=tk.LEFT, padx=5)

        tk.Label(input_frame_orders, text="Date de commande :").pack(side=tk.LEFT, padx=5)
        order_date_label = tk.Label(input_frame_orders, bg="white", width=10, relief='sunken', bd=1, cursor='hand2')
        order_date_label.pack(side=tk.LEFT, padx=5)
        order_date_label.bind('<Button-1>', GetDate)
        
        tk.Label(input_frame_orders, text="ID de client :").pack(side=tk.LEFT, padx=5)
        order_client_id_entry = tk.Entry(input_frame_orders, validate="key")
        order_client_id_entry.config(validatecommand=(frame.register(validate_id), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
        order_client_id_entry.pack(side=tk.LEFT, padx=5)
        
        tk.Label(input_frame_orders, text="Produits :").pack(side=tk.LEFT, padx=5)
        order_listprod_var = tk.StringVar(input_frame_orders, "")
        order_listprod_label = tk.Label(input_frame_orders, textvariable=order_listprod_var, bg="white", width=40, relief='sunken', bd=1, justify='left', anchor='w', cursor='hand2')
        order_listprod_label.pack(side=tk.LEFT, padx=5)
        order_listprod_label.bind('<Button-1>', ProductList)
        
        
        # Add the select boxes for Type de Transaction and Statut
        transaction_frame = tk.Frame(frame)
        transaction_frame.pack(pady=10)

        tk.Label(transaction_frame, text="Type de Transaction:").pack(side=tk.LEFT, padx=5)
        type_transaction_var = tk.StringVar(frame)
        type_transaction_var.set("Chèque")
        type_transaction_select = ttk.Combobox(transaction_frame, textvariable=type_transaction_var, values=["Chèque", "Espèces", "Virement"], state="readonly")
        type_transaction_select.pack(side=tk.LEFT, padx=5)

        tk.Label(transaction_frame, text="Statut:").pack(side=tk.LEFT, padx=5)
        statut_var = tk.StringVar(frame)
        statut_var.set("Non Payée")
        statut_select = ttk.Combobox(transaction_frame, textvariable=statut_var, values=["Payée", "Non Payée", "Avance"], state="readonly")
        statut_select.pack(side=tk.LEFT, padx=5)

        tk.Label(transaction_frame, text="Pompé :").pack(side=tk.LEFT, padx=5)
        pompe_var = tk.StringVar(frame)
        pompe_var.set("Oui")
        ttk.Combobox(transaction_frame, textvariable=pompe_var, values=["Oui", "Non"], state="readonly").pack(side=tk.LEFT, padx=5)
        
        tk.Label(transaction_frame, text="Montant payé :").pack(side=tk.LEFT, padx=5)
        order_paid_entry = tk.Entry(transaction_frame, validate="key")
        order_paid_entry.config(validatecommand=(frame.register(validate_float), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W'))
        order_paid_entry.pack(side=tk.LEFT, padx=5)
        
                        
        # Add the buttons for adding/modifying an order
        button_frame_orders = tk.Frame(frame)
        button_frame_orders.pack(pady=10)
        
        display_orders()
        
        #buttons
        ctk.CTkButton(button_frame_orders, text="Liste des commandes", command=display_orders).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame_orders, text="Ajouter Commande", command=add_order).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame_orders, text="Supprimer Commande", command=delete_order).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame_orders, text="Modifier Commande", command=modify_order).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame_orders, text="Créer un Bon de Commande", command=buttonBDC).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame_orders, text='Historique Commandes', command=get_order_history).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame_orders, text="Commandes impayées", command=get_unpaid_orders).pack(side=tk.LEFT, padx=5)
        #////////////////////////////  FIN INTERFACE ORDER ///////////////////////////////////////////
              
        
        
    ctk.CTkButton(navbar, text="Orders", command=orders).pack(side="left")
    ctk.CTkButton(navbar, text="Clients", command=Clients).pack(side="left")
    ctk.CTkButton(navbar, text="Products", command=Products).pack(side="left")
    ctk.CTkButton(navbar, text="Sales", command=Sales).pack(side="left")
    ctk.CTkButton(navbar, text="Deliveries", command=Deliveries).pack(side="left")
    ctk.CTkButton(navbar, text="Suppliers", command=Supplier).pack(side="left")


    #open_section(frame) #I had an error with this line
    window.mainloop()

# Fonction pour ouvrir différentes sections
def open_section(section_name, frame):
    Clear_widgets(frame)
    tk.Label(frame, text=section_name).pack()

# Fonction pour vider le contenu d'un cadre
def Clear_widgets(frame):
    for widget in frame.winfo_children():
        widget.destroy()

# Paramètres d'administration
ADMIN_USERNAME = "admin"
ADMIN_PASSWORD = "password"

# Fenêtre de connexion

login_window = ctk.CTk()
login_window.title("Connexion")
login_window.geometry("600x600")
login_window.configure(bg="navy")  # Set the background color to navy blue

# Create a frame for the login elements
login_frame = ctk.CTkFrame(login_window)  
login_frame.pack(pady=(50, 20), padx=20)  # Add padx to space the frame from the window edge

# Ajouter l'image
image = Image.open("./icon.png")
photo = ImageTk.PhotoImage(image)
image_label = tk.Label(login_frame, image=photo, bg="#FF6F61")  # Rouge pas foncé
image_label.grid(row=0, column=0, columnspan=2)

# Add labels and entry fields
username_label = ctk.CTkLabel(login_frame, text="Nom d'utilisateur :", font=("Helvetica", 12))  
username_label.grid(row=1, column=0, pady=(0, 5), sticky="w")

username_entry = ctk.CTkEntry(login_frame, font=("Helvetica", 12))
username_entry.grid(row=2, column=0, pady=(0, 10), sticky="w")  # Adjusted row and pady

password_label = ctk.CTkLabel(login_frame, text="Mot de passe :", font=("Helvetica", 12))  
password_label.grid(row=3, column=0, pady=(0, 5), sticky="w")

password_entry = ctk.CTkEntry(login_frame, show="*", font=("Helvetica", 12))
password_entry.grid(row=4, column=0, pady=(0, 10), sticky="w")
password_entry.bind('<Return>', login)

login_button = ctk.CTkButton(login_window, text="Se connecter", command=login, font=("Helvetica", 12))
login_button.pack(pady=(0, 20))

login_window.mainloop()